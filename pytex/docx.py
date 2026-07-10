"""DOCX backend backed by the generic reflow/document interface."""

from __future__ import annotations

import os
import math
import re
import tempfile
import uuid
import zipfile
import xml.etree.ElementTree as ET
from io import BytesIO
from pathlib import Path
from dataclasses import dataclass

from docx import Document as WordDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.table import Table as WordTable
from docx.text.paragraph import Paragraph as WordParagraph
from docx.text.run import Run as WordRun
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.shared import Pt, RGBColor, Twips
from fontTools.ttLib import TTFont, TTLibError

from pytex import align
from pytex import box as bx
from pytex.font import Font
from pytex import mmode
from pytex import node as nd
from pytex import paragraph as pg
from pytex.dimen import Dimen
from pytex.glue import Glue
from pytex.module import Module
from pytex import font_subst
from pytex import reflow
from pytex import svg


_ONE_INCH_PT = 72.0
_ONE_INCH_TEX = Dimen(72.27)
_DOCX_POINTS_PER_TEX_POINT_NUM = 7200
_DOCX_POINTS_PER_TEX_POINT_DEN = 7227
_DOCX_TWIPS_PER_TEX_POINT_NUM = 144000
_DOCX_TWIPS_PER_TEX_POINT_DEN = 7227
_DOCX_EMU_PER_TEX_POINT_NUM = 91440000
_DOCX_EMU_PER_TEX_POINT_DEN = 7227
_INLINE_TEXTBOX_PAD_PT = 0.75
_DOCX_DEFAULT_TEXT_FONT = font_subst.DEFAULT_TEXT_FONT
_MATH_FAMILY_TEXT_OVERRIDES = font_subst.MATH_FAMILY_TEXT_OVERRIDES
_MATH_OPERATORS_MAP = font_subst.MATH_OPERATORS_MAP
_MATH_LETTERS_MAP = font_subst.MATH_LETTERS_MAP
_MATH_SYMBOLS_MAP = font_subst.MATH_SYMBOLS_MAP
_MATH_LARGE_SYMBOLS_MAP = font_subst.MATH_LARGE_SYMBOLS_MAP
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
_WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
_PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"
_ASVG_NS = "http://schemas.microsoft.com/office/drawing/2016/SVG/main"
_A14_NS = "http://schemas.microsoft.com/office/drawing/2010/main"
_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
_W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
_FONT_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font"
_IMAGE_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
_OBFUSCATED_FONT_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.obfuscatedFont"
_SVG_CONTENT_TYPE = "image/svg+xml"
_WORD_FONT_FACE_ELEMENTS = {
    "regular": "embedRegular",
    "bold": "embedBold",
    "italic": "embedItalic",
    "boldItalic": "embedBoldItalic",
}
_GOTO_RE = re.compile(r"/S\s*/GoTo\b.*?/D\s*\(([^()]*)\)", re.IGNORECASE | re.DOTALL)
_GOTOR_RE = re.compile(
    r"/S\s*/GoToR\b.*?/F\s*\(([^()]*)\)(?:.*?/D\s*\(([^()]*)\))?",
    re.IGNORECASE | re.DOTALL,
)
_URI_RE = re.compile(r"/S\s*/URI\b.*?/URI\s*\(([^()]*)\)", re.IGNORECASE | re.DOTALL)
_PDF_STRING_ESCAPES = {
    "n": "\n",
    "r": "\r",
    "t": "\t",
    "b": "\b",
    "f": "\f",
    "\\": "\\",
    "(": "(",
    ")": ")",
}


ET.register_namespace("w", _W_NS)
ET.register_namespace("r", _R_NS)
ET.register_namespace("mc", _MC_NS)
ET.register_namespace("w14", _W14_NS)
ET.register_namespace("pic", _PIC_NS)
ET.register_namespace("asvg", _ASVG_NS)
ET.register_namespace("a14", _A14_NS)
ET.register_namespace("", _REL_NS)


def _color(color: reflow.Color):
    r, g, b, a = color.rgba
    return RGBColor(
        max(0, min(255, round(r * 255))),
        max(0, min(255, round(g * 255))),
        max(0, min(255, round(b * 255))),
    )


class _ContainerNode:
    def append(self, child):
        pass

def _round_docx_unit(value: float):
    if value < 0:
        return math.ceil(value - 0.5)
    return math.floor(value + 0.5)


def twips(dimen: Dimen):
    return f"{_twips(dimen)}"


def _twips(dimen: Dimen):
    return int(float(dimen) / 72.27 * 72 * 20)


def _length(dimen: Dimen):
    return Twips(_twips(dimen))


def _emu(dimen: Dimen):
    return max(1, _round_docx_unit(float(dimen) * _DOCX_EMU_PER_TEX_POINT_NUM / _DOCX_EMU_PER_TEX_POINT_DEN))


def _twip_emu(dimen: Dimen):
    return max(1, _twips(dimen) * 635)


def _docx_points(dimen: Dimen):
    return float(dimen) * _DOCX_POINTS_PER_TEX_POINT_NUM / _DOCX_POINTS_PER_TEX_POINT_DEN


def _tex_points(points: float):
    return Dimen(points * _DOCX_POINTS_PER_TEX_POINT_DEN / _DOCX_POINTS_PER_TEX_POINT_NUM)


def _svg_pt(dimen: Dimen):
    value = float(dimen) * _DOCX_POINTS_PER_TEX_POINT_NUM / _DOCX_POINTS_PER_TEX_POINT_DEN
    text = f"{value:.6f}".rstrip("0").rstrip(".")
    return f"{text or '0'}pt"


def _retarget_svg_size(payload: bytes, width: Dimen, height: Dimen):
    try:
        text = payload.decode("utf-8")
    except UnicodeDecodeError:
        return payload
    match = re.search(r"<svg\b[^>]*>", text, flags=re.IGNORECASE)
    if match is None:
        return payload

    def set_attr(tag, name, value):
        pattern = re.compile(rf'(\s{name}\s*=\s*)(["\']).*?\2', flags=re.IGNORECASE | re.DOTALL)
        replacement = rf'\1"{value}"'
        if pattern.search(tag):
            return pattern.sub(replacement, tag, count=1)
        return tag[:-1] + f' {name}="{value}">'

    tag = match.group(0)
    tag = set_attr(tag, "width", _svg_pt(width))
    tag = set_attr(tag, "height", _svg_pt(height))
    return (text[:match.start()] + tag + text[match.end():]).encode("utf-8")


def _svg_png_fallback(payload: bytes):
    try:
        import fitz
    except ImportError:
        return None
    try:
        doc = fitz.open(stream=payload, filetype="svg")
        try:
            page = doc[0]
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=True)
            return pixmap.tobytes("png")
        finally:
            doc.close()
    except Exception:
        return None


def half_pt(dimen: Dimen):
    return f"{_round_docx_unit(float(dimen) / 72.27 * 72 * 2)}"


def _textbox_xml(cx: int, cy: int, drawing_id: int):
    return f"""
<w:drawing
    xmlns:w="{_W_NS}"
    xmlns:wp="{_WP_NS}"
    xmlns:a="{_A_NS}"
    xmlns:wps="{_WPS_NS}">
  <wp:inline distT="0" distB="0" distL="0" distR="0">
    <wp:extent cx="{cx}" cy="{cy}"/>
    <wp:effectExtent l="0" t="0" r="0" b="0"/>
    <wp:docPr id="{drawing_id}" name="Inline VBox {drawing_id}"/>
    <wp:cNvGraphicFramePr>
      <a:graphicFrameLocks noChangeAspect="1"/>
    </wp:cNvGraphicFramePr>
    <a:graphic>
      <a:graphicData uri="{_WPS_NS}">
        <wps:wsp>
          <wps:cNvSpPr txBox="1"/>
          <wps:spPr>
            <a:xfrm>
              <a:off x="0" y="0"/>
              <a:ext cx="{cx}" cy="{cy}"/>
            </a:xfrm>
            <a:prstGeom prst="rect">
              <a:avLst/>
            </a:prstGeom>
            <a:noFill/>
            <a:ln>
              <a:noFill/>
            </a:ln>
          </wps:spPr>
          <wps:txbx>
            <w:txbxContent/>
          </wps:txbx>
          <wps:bodyPr wrap="none" lIns="0" tIns="0" rIns="0" bIns="0">
            <a:noAutofit/>
          </wps:bodyPr>
        </wps:wsp>
      </a:graphicData>
    </a:graphic>
  </wp:inline>
</w:drawing>
"""


def _picture_xml(
    cx: int,
    layout_cy: int,
    visual_cy: int,
    offset_y: int,
    effect_top: int,
    depth_cy: int,
    fallback_relationship_id: str,
    svg_relationship_id: str,
    drawing_id: int,
    name: str,
):
    blip_extensions = [
        f"""
                <a:ext uri="{{28A0092B-C50C-407E-A947-70E740481C1C}}">
                  <a14:useLocalDpi val="0"/>
                </a:ext>"""
    ]
    if svg_relationship_id is not None:
        blip_extensions.append(
            f"""
                <a:ext uri="{{96DAC541-7B7A-43D3-8B79-37D633B846F1}}">
                  <asvg:svgBlip r:embed="{svg_relationship_id}"/>
                </a:ext>"""
        )
    blip_extension_xml = "\n              <a:extLst>" + "".join(blip_extensions) + """
              </a:extLst>"""
    return f"""
<w:drawing
    xmlns:w="{_W_NS}"
    xmlns:r="{_R_NS}"
    xmlns:wp="{_WP_NS}"
    xmlns:a="{_A_NS}"
    xmlns:pic="{_PIC_NS}"
    xmlns:asvg="{_ASVG_NS}"
    xmlns:a14="{_A14_NS}">
  <wp:inline distT="0" distB="0" distL="0" distR="0">
    <wp:extent cx="{cx}" cy="{layout_cy}"/>
    <wp:effectExtent l="0" t="{effect_top}" r="0" b="{depth_cy}"/>
    <wp:docPr id="{drawing_id}" name="Inline Math {drawing_id}"/>
    <wp:cNvGraphicFramePr>
      <a:graphicFrameLocks noChangeAspect="1"/>
    </wp:cNvGraphicFramePr>
    <a:graphic>
      <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
        <pic:pic>
          <pic:nvPicPr>
            <pic:cNvPr id="{drawing_id}" name="{name}"/>
            <pic:cNvPicPr>
              <a:picLocks noChangeAspect="1"/>
            </pic:cNvPicPr>
          </pic:nvPicPr>
          <pic:blipFill>
            <a:blip r:embed="{fallback_relationship_id}">{blip_extension_xml}
            </a:blip>
            <a:srcRect/>
            <a:stretch>
              <a:fillRect/>
            </a:stretch>
          </pic:blipFill>
          <pic:spPr>
            <a:xfrm>
              <a:off x="0" y="{offset_y}"/>
              <a:ext cx="{cx}" cy="{visual_cy}"/>
            </a:xfrm>
            <a:prstGeom prst="rect">
              <a:avLst/>
            </a:prstGeom>
            <a:ln>
              <a:noFill/>
            </a:ln>
          </pic:spPr>
        </pic:pic>
      </a:graphicData>
    </a:graphic>
  </wp:inline>
</w:drawing>
"""


def _story_document(story):
    document = getattr(story, "document", None)
    if document is not None:
        return document
    table = getattr(story, "table", None)
    if table is not None:
        return table.document
    raise AttributeError(f"{type(story).__name__} is not attached to a DOCX document")


def _font_path(backend):
    path = getattr(backend, "path", None)
    if path:
        return path
    wrapped = getattr(backend, "_backend", None)
    return None if wrapped is None else getattr(wrapped, "path", None)


def _font_number(backend):
    wrapped = getattr(backend, "_backend", None)
    return getattr(backend, "font_number", getattr(wrapped, "font_number", 0))


def _font_kind(backend):
    return getattr(backend, "kind", getattr(getattr(backend, "_backend", None), "kind", None))


def _opentype_metric_backend(backend):
    seen = set()
    while backend is not None and id(backend) not in seen:
        seen.add(id(backend))
        if getattr(backend, "kind", None) == "opentype":
            units_per_em = getattr(backend, "units_per_em", None)
            font = getattr(backend, "font", None)
            if units_per_em and hasattr(font, "get") and font.get("hhea") is not None:
                return backend
        backend = getattr(backend, "_backend", None)
    return None


def _require_opentype_font_backend(font):
    backend = getattr(font, "backend", None)
    metric_backend = _opentype_metric_backend(backend)
    assert metric_backend is not None, (
        "DOCX output requires an OpenType-shaped font backend "
        "with kind='opentype', units_per_em, and an hhea table"
    )
    return metric_backend


@dataclass(frozen=True)
class WordFontReference:
    family: str
    face: str = "regular"

    @property
    def bold(self):
        return self.face in ("bold", "boldItalic")

    @property
    def italic(self):
        return self.face in ("italic", "boldItalic")


def _font_name_text(table, name_id, windows_only=False):
    candidates = []
    for record in getattr(table, "names", ()):
        if record.nameID != name_id:
            continue
        if windows_only and record.platformID != 3:
            continue
        try:
            text = record.toUnicode().strip()
        except Exception:
            continue
        if not text:
            continue
        candidates.append(
            (
                0 if record.platformID == 3 else 1,
                0 if record.langID in (0x0409, 0) else 1,
                text,
            )
        )
    if not candidates:
        return None
    return min(candidates)[2]


def _word_font_family(metric_backend, fallback):
    font = getattr(metric_backend, "font", None)
    table = font.get("name") if hasattr(font, "get") else None
    if table is not None:
        for name_id, windows_only in (
            (1, True),
            (16, True),
            (1, False),
            (16, False),
            (4, True),
            (6, True),
            (4, False),
            (6, False),
        ):
            family = _font_name_text(table, name_id, windows_only=windows_only)
            if family:
                return family
    return fallback


def _word_font_face(metric_backend):
    font = getattr(metric_backend, "font", None)
    if not hasattr(font, "get"):
        return "regular"

    table = font.get("name")
    subfamily = None
    if table is not None:
        for name_id, windows_only in ((2, True), (17, True), (2, False), (17, False)):
            subfamily = _font_name_text(table, name_id, windows_only=windows_only)
            if subfamily:
                break
    style = re.sub(r"[^a-z0-9]+", "", (subfamily or "").casefold())
    bold = any(term in style for term in ("bold", "semibold", "demibold"))
    italic = any(term in style for term in ("italic", "oblique", "slanted"))

    os2 = font.get("OS/2")
    if os2 is not None:
        selection = int(getattr(os2, "fsSelection", 0))
        bold = bold or bool(selection & (1 << 5)) or int(getattr(os2, "usWeightClass", 0)) >= 600
        italic = italic or bool(selection & ((1 << 0) | (1 << 9)))
    head = font.get("head")
    if head is not None:
        mac_style = int(getattr(head, "macStyle", 0))
        bold = bold or bool(mac_style & 1)
        italic = italic or bool(mac_style & 2)
    post = font.get("post")
    if post is not None:
        italic = italic or float(getattr(post, "italicAngle", 0)) != 0

    if bold and italic:
        return "boldItalic"
    if bold:
        return "bold"
    if italic:
        return "italic"
    return "regular"


def _docx_font_reference(backend):
    fallback = font_subst.fontBackendName(backend) or getattr(backend, "name", None)
    metric_backend = _opentype_metric_backend(backend)
    if metric_backend is None:
        return None if fallback is None else WordFontReference(fallback)
    family = _word_font_family(metric_backend, fallback)
    if not family:
        return None
    return WordFontReference(family, _word_font_face(metric_backend))


def _docx_font_name(backend):
    reference = _docx_font_reference(backend)
    return None if reference is None else reference.family


def _font_key_bytes(font_key):
    data = bytes.fromhex(re.sub(r"[{}-]", "", font_key))
    guid_memory = data[3::-1] + data[5:3:-1] + data[7:5:-1] + data[8:16]
    return guid_memory[::-1]


def _obfuscate_font(data, font_key):
    out = bytearray(data)
    key = _font_key_bytes(font_key)
    for index in range(min(32, len(out))):
        out[index] ^= key[index % 16]
    return bytes(out)


def _docx_font_payload(path, font_number=0):
    """Return embeddable font bytes and their obfuscated-part suffix.

    Word only reliably embeds TrueType outlines. Convert a selected CFF face
    to a standalone TrueType font here, at the DOCX packaging boundary, while
    leaving existing TrueType fonts and collections untouched.
    """
    with open(path, "rb") as font_file:
        data = font_file.read()
    source_suffix = (
        ".odttc"
        if Path(path).suffix.lower() in {".ttc", ".otc"}
        else ".odttf"
    )
    try:
        font = TTFont(
            BytesIO(data),
            fontNumber=font_number,
            lazy=False,
            recalcTimestamp=False,
        )
    except (OSError, TTLibError):
        return data, source_suffix
    try:
        if "CFF " not in font:
            return data, source_suffix
        from afdko.otf2ttf import otf_to_ttf

        otf_to_ttf(font)
        converted = BytesIO()
        font.save(converted)
        return converted.getvalue(), ".odttf"
    finally:
        font.close()


def _xml_bytes(node):
    return b"<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n" + ET.tostring(node, encoding="utf-8")


def _decode_pdf_string(token):
    if len(token) < 2 or token[0] != "(" or token[-1] != ")":
        return token
    out = []
    i = 1
    while i < len(token) - 1:
        ch = token[i]
        if ch == "\\" and i + 1 < len(token) - 1:
            i += 1
            esc = token[i]
            out.append(_PDF_STRING_ESCAPES.get(esc, esc))
        else:
            out.append(ch)
        i += 1
    return "".join(out)


def _annotation_info(payload):
    goto = _GOTO_RE.search(payload)
    if goto is not None:
        return {
            "kind": "goto",
            "destination": _decode_pdf_string(f"({goto.group(1)})"),
        }
    gotor = _GOTOR_RE.search(payload)
    if gotor is not None:
        return {
            "kind": "gotor",
            "file": _decode_pdf_string(f"({gotor.group(1)})"),
            "destination": None if gotor.group(2) is None else _decode_pdf_string(f"({gotor.group(2)})"),
        }
    uri = _URI_RE.search(payload)
    if uri is not None:
        return {
            "kind": "uri",
            "url": _decode_pdf_string(f"({uri.group(1)})"),
        }
    return {
        "kind": "raw",
        "payload": payload,
    }


class Text(reflow.Element):
    XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

    def __init__(self, preserve_space=False):
        node = OxmlElement("w:t")
        if preserve_space:
            node.set(self.XML_SPACE, "preserve")
        node.text = ""
        super().__init__(node)

    def setChar(self, char: nd.Node):
        if char.node_type == nd.NODE_TYPE.LIGATURE:
            for child in char.source:
                self.setChar(child)
        else:
            self._node.text += char.char


class TextBoxStory(reflow.Block):
    def __init__(self, document, drawing, node, box: bx.Box):
        super().__init__(node, inline=True)
        self.document = document
        self.drawing = drawing
        self.box = box

    @property
    def line_id(self):
        return self.document.line_id

    @property
    def part(self):
        return self.document._node.part

    def _new_word_paragraph(self):
        node = OxmlElement("w:p")
        self._node.append(node)
        return WordParagraph(node, self)

    def _new_word_table(self):
        node = CT_Tbl.new_tbl(0, 0, Twips(0))
        self._node.append(node)
        return WordTable(node, self)

    def newParagraph(self, spacing_before=Dimen(), justify: str = "left") -> "Paragraph":
        para = Paragraph(self, spacing_before=spacing_before, justify=justify)
        self.nodes.append(para)
        return para

    def newTable(self, xspacing=Dimen(), yspacing=Dimen()):
        table = Table(
            self.document,
            self._new_word_table(),
            xspacing=xspacing,
            yspacing=yspacing,
            full_width=self.box.width,
        )
        self.nodes.append(table)
        return table

    def newGraph(self, key, type, file):
        return None

    def finalizeContent(self):
        for node in self.nodes:
            if isinstance(node, Table):
                node.setFullWidth(self.box.width)


class TextRun(reflow.TextRun):
    def __init__(
        self,
        line,
        text="",
        font=None,
        color=reflow.Color.black,
        baseline_from_bottom=Dimen(),
        preserve_space=False,
        ):
        node = line._node.add_run()
        self.line = line
        self.has_text_glyphs = False
        super().__init__(
            node,
            text=text,
            font=font,
            color=color,
            baseline_from_bottom=baseline_from_bottom,
        )
        self.preserve_space = preserve_space
        if text:
            self.line.applyLeadingSpacing()
            self.has_text_glyphs = True
            self.uses_backend_baseline = True
        t = self.newText()
        t._node.text = "" if text is None else text
        rPr = node._r.get_or_add_rPr()
        kern = OxmlElement("w:kern")
        rPr.append(kern)
        kern.set(qn("w:val"), "1")
        lig = OxmlElement("w14:ligatures")
        rPr.append(lig)
        lig.set(qn("w14:val"), "standard")
        self.setFont(font)

    def setFont(self, font):
        self.font = font
        self.line.font = font
        if font is not None:
            document = _story_document(self.line.story)
            reference = document.defineFont(font)
            if reference is not None:
                self._node.font.name = reference.family
                rPr = self._node._r.get_or_add_rPr()
                rFonts = rPr.rFonts
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.insert(0, rFonts)
                for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
                    rFonts.set(qn(f"w:{attr}"), reference.family)
                if reference.bold:
                    self._node.font.bold = True
                if reference.italic:
                    self._node.font.italic = True
            self._node.font.size = Pt(round(float(font.at) / 72.27 * 72 * 2) / 2)
            self._node.font.color.rgb = _color(self.color)

    def newText(self) -> Text:
        text = Text(self.preserve_space)
        self._node._element.append(text._node)
        self.nodes.append(text)
        self.text = text
        return text

    def setChar(self, char: nd.Node):
        self.line.applyLeadingSpacing()
        self.has_text_glyphs = True
        self.uses_backend_baseline = True
        if self.text is None:
            self.newText()
        self.text.setChar(char)

    def newSpace(self, width: Dimen, breakable: bool):
        if not self.line.has_visible_content:
            self.line.leading_spacing += width
            return None
        self.line.applyLeadingSpacing()
        if self.text is None:
            self.newText()
        self.text._node.text = " " if breakable else "\xa0"
        self.text._node.set(Text.XML_SPACE, "preserve")
        self.uses_backend_baseline = True
        if self.font is None:
            return self
        diff = width - self.font.at * self.font.backend._spaceWidth()
        if int(diff) != 0:
            rPr = self._node._r.get_or_add_rPr()
            spacing_element = OxmlElement("w:spacing")
            spacing_element.set(qn("w:val"), twips(diff))
            rPr.append(spacing_element)
        return self

    def verticalShift(self, shift):
        self._setPosition(_round_docx_unit(float(shift) * 2))

    def _setPosition(self, position):
        if not position:
            return
        rPr = self._node._r.get_or_add_rPr()
        existing = rPr.find(qn("w:position"))
        if existing is not None:
            rPr.remove(existing)
        position_element = OxmlElement("w:position")
        position_element.set(qn("w:val"), str(position))
        rPr.append(position_element)

    def _boxBaselineFromBottom(self, depth):
        return self.baseline_from_bottom - self.line.baseline_from_bottom + depth

    def newInlineVBox(self, box: bx.Box):
        self.line.applyLeadingSpacing()
        self.text = None
        self.uses_backend_baseline = False
        document = _story_document(self.line.story)
        drawing_id = document.nextDrawingId()
        cy = _emu(box.height + box.depth)
        drawing = parse_xml(
            _textbox_xml(
                _emu(box.width),
                cy,
                drawing_id,
            )
        )
        content = drawing.find(f".//{{{_W_NS}}}txbxContent")
        if content is None:
            raise ValueError("DOCX inline textbox template is missing w:txbxContent")
        block = TextBoxStory(document, drawing, content, box)
        self._node._element.append(drawing)
        self.baseline_from_bottom = self._boxBaselineFromBottom(box.depth)
        self.nodes.append(block)
        return block

    def newInlineMath(self, backend, inlinemath: mmode.InlineMathNode, box: bx.Box, piece: int):
        self.line.applyLeadingSpacing()
        self.text = None
        self.uses_backend_baseline = False
        document = _story_document(self.line.story)
        payload = backend.inlineMathSvg(box)
        visual_height = box.height + box.depth
        is_display_math = isinstance(box, DisplayMathPictureBox)
        if is_display_math:
            payload = _retarget_svg_size(payload, box.width, visual_height)
        fallback_placeholder, svg_placeholder, media_name = document.defineInlineSvg(
            payload,
            width=box.width,
            height=visual_height,
            use_svg=not is_display_math,
        )
        drawing_id = document.nextDrawingId()
        drawing = parse_xml(
            _picture_xml(
                _emu(box.width),
                _emu(visual_height),
                _emu(visual_height),
                0,
                0,
                0,
                fallback_placeholder,
                svg_placeholder,
                drawing_id,
                media_name,
            )
        )
        self.baseline_from_bottom = self._boxBaselineFromBottom(box.depth)
        self._node._element.append(drawing)
        self.line.addInlineDrawing(drawing)
        picture = reflow.Element(drawing)
        self.nodes.append(picture)
        return picture

    def newInlineGraphic(self, backend, asset, request):
        width = request.width or asset.width
        height = request.height or asset.height
        if width is None or height is None:
            return None
        payload = backend.graphicSvgPayload(asset, width, height)
        if payload is None:
            return None
        visual_height = height + asset.depth
        standalone_graphic_line = not self.line.has_visible_content
        self.line.applyLeadingSpacing()
        self.text = None
        self.uses_backend_baseline = False
        document = _story_document(self.line.story)
        fallback_placeholder, svg_placeholder, media_name = document.defineInlineSvg(
            payload,
            width=width,
            height=visual_height,
        )
        drawing_id = document.nextDrawingId()
        visual_cy = _twip_emu(visual_height) if standalone_graphic_line else _emu(visual_height)
        drawing = parse_xml(
            _picture_xml(
                _emu(width),
                visual_cy,
                visual_cy,
                0,
                0,
                0,
                fallback_placeholder,
                svg_placeholder,
                drawing_id,
                media_name,
            )
        )
        self.baseline_from_bottom = self._boxBaselineFromBottom(asset.depth)
        self._node._element.append(drawing)
        self.line.addInlineDrawing(drawing)
        picture = reflow.Element(drawing)
        self.nodes.append(picture)
        return picture

class Space(TextRun):
    def __init__(self, line, width: Dimen, breakable: bool, font: Font):
        space = " " if breakable else "\xa0"
        super().__init__(
            line,
            space,
            font,
            baseline_from_bottom=line.baseline_from_bottom,
            preserve_space=True,
        )
        self.has_text_glyphs = False
        self.uses_backend_baseline = True
        diff = width - font.at * font.backend._spaceWidth()
        if int(diff) != 0:
            rPr = self._node._r.get_or_add_rPr()
            spacing_element = OxmlElement('w:spacing')
            spacing_element.set(qn('w:val'), twips(diff))
            rPr.append(spacing_element)


class Line(reflow.Line):
    JUSTIFY = {
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        None: WD_ALIGN_PARAGRAPH.LEFT,
    }
    drop_trailing_breakable_spacing = True

    def __init__(
        self,
        para: WordParagraph,
        line_id: int,
        line_spec: reflow.LineSpec,
        justify="justify",
        story=None,
    ):
        super().__init__(para, line_spec)
        self.line_spec = line_spec
        self.story = story
        self.justify = self._wordJustify(justify)
        para.alignment = self.justify
        self.line_height = line_spec.line_height
        self.inline_drawings = []
        self.leading_spacing = Dimen()
        self.has_visible_content = False
        self.has_text_glyphs = False
        self._setLineHeight(self.line_height)
        fmt = para.paragraph_format
        self.spacing_before = Dimen(line_spec.spacing_before)
        self._setSpaceBefore(self.spacing_before)
        fmt.space_after = Pt(0)
        self.font = line_spec.default_font
        self.width = line_spec.line_box.rightmost()
        self.line_id = line_id

    def _wordJustify(self, justify):
        if isinstance(self.story, Story):
            return WD_ALIGN_PARAGRAPH.JUSTIFY
        return self.JUSTIFY[justify]

    def _setLineHeight(self, height):
        fmt = self._node.paragraph_format
        fmt.line_spacing = Twips(max(1, _twips(height)))
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    def _setSpaceBefore(self, spacing):
        pPr = self._node._p.get_or_add_pPr()
        spacing_element = pPr.find(qn("w:spacing"))
        if spacing_element is None:
            spacing_element = OxmlElement("w:spacing")
            pPr.append(spacing_element)
        spacing_element.set(qn("w:before"), str(_twips(spacing)))

    def setLineHeight(self, height):
        self.line_height = Dimen(height)
        self._setLineHeight(self.line_height)

    @property
    def node(self):
        return self._node

    def registerTextRun(self, run: TextRun):
        self.nodes.append(run)
        if getattr(run, "has_text_glyphs", False):
            self.has_text_glyphs = True
        return run

    def backendBaselineForFont(self, font):
        if font is None:
            return None
        backend = _require_opentype_font_backend(font)
        font_size = round(_docx_points(font.at) * 2) / 2
        line_height = _docx_points(self.line_height)
        line_baseline = getattr(backend, "lineBaselineFromBottom", None)
        if callable(line_baseline):
            try:
                baseline = line_baseline(font_size, line_height)
            except TypeError:
                baseline = line_baseline(font_size, line_height, round_total=None)
            baseline = _docx_points(Dimen(baseline))
            return baseline if baseline > 0 else None
        hhea = backend.font.get("hhea")
        units_per_em = backend.units_per_em
        ascent = max(0, getattr(hhea, "ascent", 0))
        descent = max(0, -getattr(hhea, "descent", 0))
        line_gap = max(0, getattr(hhea, "lineGap", 0))
        # Match Word's natural font box, then scale the baseline to exact line height.
        padding = round(0.15 * (ascent + descent))
        total_units = ascent + descent + line_gap + 2 * padding
        if total_units <= 0:
            return None
        total_size = math.ceil(total_units / units_per_em * font_size * 2 - 1e-9) / 2
        if total_size <= 0:
            return None
        baseline = _docx_points(self.line_height) * (
            (descent + padding) / units_per_em * font_size
        ) / total_size
        return baseline if baseline > 0 else None

    def finalizeLine(self):
        pass

    def addInlineDrawing(self, drawing):
        self.inline_drawings.append(drawing)

    def applyTrailingSpacing(self, spacing: Dimen):
        pass

    def _shrinkLastDrawingLayout(self, reduction: Dimen):
        if not self.inline_drawings:
            return
        drawing = self.inline_drawings[-1]
        extent = drawing.find(f".//{{{_WP_NS}}}extent")
        effect = drawing.find(f".//{{{_WP_NS}}}effectExtent")
        if extent is None:
            return
        old_cy = int(extent.get("cy", "1"))
        new_cy = max(1, old_cy - _emu(reduction))
        actual_reduction = old_cy - new_cy
        extent.set("cy", str(new_cy))
        if effect is not None and actual_reduction > 0:
            bottom = int(effect.get("b", "0"))
            effect.set("b", str(bottom + actual_reduction))

    def newTextRun(self, text, font, color, baseline_from_bottom) -> TextRun:
        self.registerBackendBaseline(font)
        return self.registerTextRun(
            TextRun(
                self,
                text=text,
                font=font,
                color=color,
                baseline_from_bottom=baseline_from_bottom,
            )
        )

    def newSpace(self, width: Dimen, breakable: bool):
        if not self.has_visible_content:
            self.leading_spacing += width
            return None
        s = Space(self, width, breakable, self.font)
        self.registerTextRun(s)
        return s

    def applyLeadingSpacing(self):
        if self.has_visible_content:
            return
        if int(self.leading_spacing) > 0:
            self._node.paragraph_format.left_indent = Twips(_twips(self.leading_spacing))
        self.leading_spacing = Dimen()
        self.has_visible_content = True


class HyperlinkRunContainer:
    def __init__(self, paragraph: WordParagraph, element):
        self.paragraph = paragraph
        self._element = element

    @property
    def part(self):
        return self.paragraph.part

    def add_run(self):
        run = OxmlElement("w:r")
        self._element.append(run)
        return WordRun(run, self.paragraph)


class AnnotationLine(reflow.Line):
    def __init__(self, parent: Line, element):
        self.parent = parent
        self.story = parent.story
        self._node = HyperlinkRunContainer(parent._node, element)
        self.nodes = []
        self.line_spec = parent.line_spec
        self.font = parent.font
        self.baseline_from_bottom = parent.baseline_from_bottom
        self.lign_height = getattr(parent, "lign_height", None)
        self.line_height = parent.line_height
        self.leading_spacing = Dimen()
        self.has_visible_content = False
        self.backend_baseline = parent.backend_baseline

    def backendBaselineForFont(self, font):
        return self.parent.backendBaselineForFont(font)

    def registerBackendBaseline(self, font):
        self.parent.registerBackendBaseline(font)
        self.backend_baseline = self.parent.backend_baseline

    def newTextRun(self, text, font, color, baseline_from_bottom) -> TextRun:
        self.registerBackendBaseline(font)
        run = TextRun(
            self,
            text=text,
            font=font,
            color=color,
            baseline_from_bottom=baseline_from_bottom,
        )
        self.nodes.append(run)
        self.parent.registerTextRun(run)
        return run

    def newSpace(self, width: Dimen, breakable: bool):
        if not self.has_visible_content and not self.parent.has_visible_content:
            self.parent.newSpace(width, breakable)
            return None
        run = Space(self, width, breakable, self.font)
        self.nodes.append(run)
        self.parent.registerTextRun(run)
        return run

    def applyLeadingSpacing(self):
        if not self.parent.has_visible_content:
            self.parent.applyLeadingSpacing()
        self.has_visible_content = True

    def addInlineDrawing(self, drawing):
        self.parent.addInlineDrawing(drawing)

    def applyTrailingSpacing(self, spacing: Dimen):
        self.parent.applyTrailingSpacing(spacing)


class AnnotationBuilder(reflow.AnnotationBuilder):
    def __init__(self, backend, parent, name, anchor=None, href=None):
        super().__init__(backend, parent, name)
        self.anchor = anchor
        self.href = href
        self.element = None

    def beginAnnotation(self, name):
        line = self.parent.container
        if not self.href and not self.anchor:
            self.container = line
            return
        hyperlink = OxmlElement("w:hyperlink")
        if self.href:
            rid = line._node.part.relate_to(self.href, RT.HYPERLINK, is_external=True)
            hyperlink.set(qn("r:id"), rid)
        elif self.anchor:
            hyperlink.set(qn("w:anchor"), self.anchor)
        line._node._p.append(hyperlink)
        self.element = hyperlink
        self.container = AnnotationLine(line, hyperlink)

    def endAnnotation(self, name):
        pass


class Paragraph(reflow.Paragraph):
    def __init__(self, story, spacing_before=Dimen(), justify="justify"):
        node = _ContainerNode()
        super().__init__(node, spacing_before, justify)
        self.story = story
        self.spacing = spacing_before

    def setJustify(self, justify):
        self.justify = justify

    def newLine(self, line_spec: reflow.LineSpec) -> Line:
        para = self.story._new_word_paragraph()
        if self.spacing is not None:
            line_spec.spacing_before += self.spacing
            self.spacing = None
        line = Line(para, self.story.line_id, line_spec, justify=self.justify, story=self.story)
        self.append(line)
        return line

    def applyTrailingSpacing(self, spacing: Dimen):
        if self.nodes:
            self.nodes[-1].applyTrailingSpacing(spacing)
        elif self.spacing is not None:
            self.spacing += spacing


class Cell(reflow.Cell):
    def __init__(self, row, node, span=1, width=None, justify: str = "justify"):
        super().__init__(node, span=span, width=width, justify=justify)
        self.row = row
        self.table = row.table
        self._used_initial_paragraph = False
        self._used_row_spacing = False
        node.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        self._setWidth(width)
        self._setNoWrap()

    @property
    def line_id(self):
        return self.table.line_id

    def _new_word_paragraph(self):
        if not self._used_initial_paragraph:
            self._used_initial_paragraph = True
            paragraphs = self._node.paragraphs
            if paragraphs and not paragraphs[0].text and not paragraphs[0].runs:
                return paragraphs[0]
        return self._node.add_paragraph()

    def _setWidth(self, width):
        if width is None:
            return
        tcPr = self._node._tc.get_or_add_tcPr()
        tcW = tcPr.tcW
        if tcW is None:
            tcW = OxmlElement("w:tcW")
            tcPr.append(tcW)
        if isinstance(width, Dimen):
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), twips(width))
            return
        if isinstance(width, (int, float)):
            tcW.set(qn("w:type"), "pct")
            tcW.set(qn("w:w"), str(int(float(width) * 5000)))

    def _setNoWrap(self):
        tcPr = self._node._tc.get_or_add_tcPr()
        if tcPr.find(qn("w:noWrap")) is None:
            tcPr.append(OxmlElement("w:noWrap"))

    def newParagraph(self) -> Paragraph:
        spacing_before = Dimen()
        if not self._used_row_spacing:
            spacing_before = self.row.spacing_before
            self._used_row_spacing = True
        para = Paragraph(self, spacing_before=spacing_before, justify=self.justify)
        self.nodes.append(para)
        return para


class Row(reflow.Row):
    def __init__(self, table, node, row_box=None, spacing_before=Dimen()):
        super().__init__(node)
        self.table = table
        self._cell_index = 0
        self.spacing_before = Dimen(spacing_before)
        self.row_box = row_box
        self._setHeight(row_box, self.spacing_before)

    def _setHeight(self, row_box, spacing_before):
        if row_box is None:
            return
        height = row_box.height + row_box.depth + spacing_before
        self._node.height = Twips(max(1, _twips(height)))
        self._node.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    def newCell(self, span=1, width=None, relative_width=None, justify="justify") -> Cell:
        node = self.table._wordCell(self._node, self._cell_index, span, width)
        self._cell_index += span
        cell = Cell(self, node, span=span, width=width, justify=justify)
        self.nodes.append(cell)
        return cell


class Table(reflow.Table):
    ALIGNMENT = {
        "left": WD_TABLE_ALIGNMENT.LEFT,
        "center": WD_TABLE_ALIGNMENT.CENTER,
        "right": WD_TABLE_ALIGNMENT.RIGHT,
    }

    def __init__(
        self,
        document,
        node,
        xspacing=Dimen(),
        yspacing=Dimen(),
        full_width=None,
        alignment="left",
    ):
        super().__init__(node, xspacing=xspacing, yspacing=yspacing)
        self.document = document
        self.owner = None
        self.box = None
        self.space_before = Dimen(yspacing)
        self.region = "body"
        self.full_width = None if full_width is None else Dimen(full_width)
        self.alignment = alignment
        self._node.autofit = self.full_width is None
        self._setAlignment()
        self._setCellMargins()
        if self.full_width is not None:
            self.setFullWidth(self.full_width)

    @property
    def line_id(self):
        return self.document.line_id

    def _setAlignment(self):
        self._node.alignment = self.ALIGNMENT.get(self.alignment, WD_TABLE_ALIGNMENT.LEFT)

    def _setCellMargins(self):
        tblPr = self._node._tbl.tblPr
        cellMar = tblPr.first_child_found_in("w:tblCellMar")
        if cellMar is None:
            cellMar = OxmlElement("w:tblCellMar")
            tblPr.append(cellMar)
        for side in ("top", "left", "bottom", "right"):
            existing = cellMar.find(qn(f"w:{side}"))
            if existing is not None:
                cellMar.remove(existing)
            margin = OxmlElement(f"w:{side}")
            margin.set(qn("w:w"), "0")
            margin.set(qn("w:type"), "dxa")
            cellMar.append(margin)

    def _setTableWidth(self, width_type, width):
        tblPr = self._node._tbl.tblPr
        tblW = tblPr.first_child_found_in("w:tblW")
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.insert(0, tblW)
        tblW.set(qn("w:type"), width_type)
        tblW.set(qn("w:w"), str(width))

    def _setTableLayoutFixed(self):
        tblPr = self._node._tbl.tblPr
        tblLayout = tblPr.first_child_found_in("w:tblLayout")
        if tblLayout is None:
            tblLayout = OxmlElement("w:tblLayout")
            tblPr.append(tblLayout)
        tblLayout.set(qn("w:type"), "fixed")

    @staticmethod
    def _columnWidth(width=None):
        if isinstance(width, Dimen):
            return Twips(max(1, _twips(width)))
        if isinstance(width, (int, float)):
            return Twips(max(1, int(1440 * float(width))))
        return Twips(1)

    def _setColumnWidth(self, index, width):
        if width is None:
            return
        width = self._columnWidth(width)
        column = self._node.columns[index]
        if column.width is None or int(column.width) < int(width):
            column.width = width

    def _ensureColumns(self, count, width=None):
        while len(self._node.columns) < count:
            self._node.add_column(self._columnWidth(width))
        if count:
            self._setTableLayoutFixed()

    def _wordCell(self, row, index, span=1, width=None):
        span = max(1, int(span))
        column_width = width
        if isinstance(width, Dimen) and span > 1:
            column_width = width / span
        self._ensureColumns(index + span, column_width)
        for column_index in range(index, index + span):
            self._setColumnWidth(column_index, column_width)
        cell = row.cells[index]
        if span > 1:
            cell = cell.merge(row.cells[index + span - 1])
        return cell

    def setFullWidth(self, width):
        self.full_width = Dimen(width)
        self._node.autofit = False
        self._setTableWidth("dxa", twips(self.full_width))
        self._setTableLayoutFixed()
        grid = self._node._tbl.tblGrid
        columns = list(grid.gridCol_lst)
        if not columns:
            return
        current = []
        for column in columns:
            value = column.get(qn("w:w"))
            current.append(max(1, int(value)) if value is not None else 1)
        total = sum(current)
        if total <= 0:
            current = [1] * len(columns)
        for column, value in zip(columns, current):
            column.set(qn("w:w"), str(value))
        self._setCellWidths(current)

    def _setCellWidths(self, columns):
        for row in self._node._tbl.tr_lst:
            column_index = 0
            for cell in row.tc_lst:
                tcPr = cell.get_or_add_tcPr()
                grid_span = tcPr.gridSpan
                span = int(grid_span.val) if grid_span is not None else 1
                width = sum(columns[column_index:column_index + span])
                tcW = tcPr.tcW
                if tcW is None:
                    tcW = OxmlElement("w:tcW")
                    tcPr.append(tcW)
                tcW.set(qn("w:type"), "dxa")
                tcW.set(qn("w:w"), str(max(1, width)))
                column_index += span

    def newRow(self, row_box=None, spacing_before=Dimen()) -> Row:
        row = Row(self, self._node.add_row(), row_box=row_box, spacing_before=spacing_before)
        self.nodes.append(row)
        return row

    def iter_specs(self):
        yield self


class Block(reflow.Block):
    def __init__(self, backend, region="body", inline=False, xspacing=Dimen(), yspacing=Dimen()):
        super().__init__(_ContainerNode(), inline=inline, xspacing=xspacing, yspacing=yspacing)
        self.backend = backend
        self.region = region
        self._entries = []

    def newParagraph(self, spacing_before=Dimen(), justify: str = "left") -> Paragraph:
        para = Paragraph(self._node, spacing_before=spacing_before, justify=justify)
        self.append(para)
        return para

    def newTable(self, xspacing=Dimen(), yspacing=Dimen()):
        raise NotImplementedError("DOCX inline block tables need a real story container")

    def newGraph(self, key, type, file):
        return None


class Story(reflow.Element):
    def __init__(self, document, node, section=None, region="body"):
        super().__init__(node)
        self.document = document
        self.section = section
        self.region = region

    def clear(self):
        element = getattr(self._node, "_element", None)
        if element is None:
            return
        for child in list(element):
            if child.tag in (qn("w:p"), qn("w:tbl")):
                element.remove(child)
        self.nodes.clear()

    def newParagraph(self, spacing_before=Dimen(), justify: str = "left") -> Paragraph:
        para = Paragraph(self, spacing_before=spacing_before, justify=justify)
        self.nodes.append(para)
        return para

    def _new_word_paragraph(self):
        return self._node.add_paragraph()

    def _new_word_table(self):
        try:
            return self._node.add_table(rows=0, cols=0, width=Twips(0))
        except TypeError:
            return self._node.add_table(rows=0, cols=0)

    def newTable(self, xspacing=Dimen(), yspacing=Dimen()):
        table = Table(self.document, self._new_word_table(), xspacing=xspacing, yspacing=yspacing)
        self.nodes.append(table)
        return table

    def newGraph(self, key, type, file):
        return None
    
    @property
    def line_id(self):
        return self.document.line_id

    def applyTrailingSpacing(self, spacing: Dimen):
        if self.region == "body" and self.section is not None:
            self.section.applyTrailingSpacing(spacing)
            return
        if self.nodes and hasattr(self.nodes[-1], "applyTrailingSpacing"):
            self.nodes[-1].applyTrailingSpacing(spacing)


class Section:
    def __init__(self, document, spec: reflow.PageSpec):
        self.document = document
        self.spec = spec
        self._section = document._node.sections[-1]
        self._apply_spec()
        self._section.header.is_linked_to_previous = False
        self._section.footer.is_linked_to_previous = False
        self._header = Story(document, self._section.header, self, "header")
        self._footer = Story(document, self._section.footer, self, "footer")
        self._body = Story(document, document._node._body, self, "body")

    def _apply_spec(self):
        section = self._section
        section.page_width = _length(self.spec.width)
        section.page_height = _length(self.spec.height)
        section.left_margin = _length(self.spec.margin_left)
        section.top_margin = _length(self.spec.margin_top)
        section.right_margin = _length(self.spec.margin_right)
        section.bottom_margin = _length(self.spec.margin_bottom)
        if self.spec.header_distance is not None:
            section.header_distance = _length(self.spec.header_distance)
        if self.spec.footer_distance is not None:
            section.footer_distance = _length(self.spec.footer_distance)

    def applyTrailingSpacing(self, spacing: Dimen):
        if int(spacing) >= 0:
            return
        current = self._section.bottom_margin
        if current is None:
            return
        reduction = _twips(Dimen() - spacing)
        self._section.bottom_margin = Twips(max(0, int(current.twips) - reduction))

    @property
    def header(self) -> Block:
        return self._header

    @property
    def body(self) -> Block:
        return self._body

    @property
    def footer(self) -> Block:
        return self._footer

    def close(self, document, last_page):
        document.add_section(WD_SECTION_START.NEW_PAGE)
        self._minimize_section_break_paragraph(document)

    @staticmethod
    def _minimize_section_break_paragraph(document):
        """Make python-docx's empty section-break paragraph take minimal space."""
        paragraphs = document.paragraphs
        if not paragraphs:
            return
        paragraph = paragraphs[-1]
        pPr = paragraph._p.pPr
        if pPr is None or pPr.find(qn("w:sectPr")) is None:
            return
        fmt = paragraph.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = Twips(1)
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY


@dataclass
class EmbeddedFont:
    family: str
    face: str
    path: str
    font_number: int = 0


@dataclass
class InlineSvgPicture:
    svg_placeholder: str
    media_name: str
    payload: bytes
    fallback_placeholder: str = None
    fallback_media_name: str = None
    fallback_payload: bytes = None


class DisplayMathPictureBox:
    node_type = nd.NODE_TYPE.VLIST

    def __init__(self, box: bx.Box):
        self.list = [box]
        self.source = getattr(box, "source", None)
        shifted = Dimen(getattr(box, "shifted", Dimen()))
        self.width = Dimen(getattr(box, "width", Dimen()))
        if shifted > 0:
            self.width += shifted
        self.height = Dimen(getattr(box, "height", Dimen()))
        self.depth = Dimen(getattr(box, "depth", Dimen()))
        self.to = self.height
        self.spread = Dimen()
        self.natural = Glue()
        self.glue_ratio = bx.GlueRatio(0, 0, 1)
        self.shifted = Dimen()

    def rightmost(self):
        return self.width


class Document(reflow.Document):
    def __init__(self, title: str, output=None):
        document = WordDocument()
        super().__init__(document, title, output)
        self.sections = []
        self._line_id = 0
        self._drawing_id = 0
        self._bookmark_id = 0
        self._embedded_fonts = {}
        self._inline_svg_pictures = {}

    @property
    def line_id(self):
        self._line_id += 1
        return self._line_id

    def nextDrawingId(self):
        self._drawing_id += 1
        return self._drawing_id

    def nextBookmarkId(self):
        bookmark_id = self._bookmark_id
        self._bookmark_id += 1
        return bookmark_id

    @property
    def header(self) -> Block:
        return self.sections[-1].header

    @property
    def body(self) -> Block:
        return self.sections[-1].body

    @property
    def footer(self) -> Block:
        return self.sections[-1].footer

    def newPage(self, page_spec: reflow.PageSpec) -> Section:
        """Start a new DOCX section for each shipped TeX page."""
        section_index = len(self.sections)
        if section_index > 0:
            self.sections[-1].close(self._node, last_page=False)
        section = Section(self, page_spec)
        self.sections.append(section)
        return section

    def defineFont(self, font):
        if font is None:
            return None
        backend = getattr(font, "backend", None)
        if backend is None:
            return None
        _require_opentype_font_backend(font)
        reference = _docx_font_reference(backend)
        if reference is None:
            return None
        path = _font_path(backend)
        if _font_kind(backend) == "opentype" and isinstance(path, str) and os.path.isfile(path):
            path = os.path.realpath(path)
            key = (reference.family, reference.face)
            if key not in self._embedded_fonts:
                self._embedded_fonts[key] = EmbeddedFont(
                    reference.family,
                    reference.face,
                    path,
                    _font_number(backend),
                )
        return reference

    def definePicture(self, key, type, path):
        return None

    def defineInlineSvg(self, payload: bytes, width=None, height=None, use_svg=True):
        index = len(self._inline_svg_pictures) + 1
        svg_placeholder = f"pytexInlineSvg{index}"
        fallback_placeholder = f"pytexInlinePng{index}"
        media_name = f"pytex-inline-math-{index}.svg"
        fallback_media_name = f"pytex-inline-math-{index}.png"
        fallback_payload = _svg_png_fallback(payload)
        if not use_svg and fallback_payload is not None:
            svg_placeholder = None
            media_name = None
        key = svg_placeholder or fallback_placeholder
        self._inline_svg_pictures[key] = InlineSvgPicture(
            svg_placeholder,
            media_name,
            payload,
            fallback_placeholder=fallback_placeholder if fallback_payload is not None else svg_placeholder,
            fallback_media_name=fallback_media_name if fallback_payload is not None else media_name,
            fallback_payload=fallback_payload if fallback_payload is not None else payload,
        )
        fallback_reference = fallback_placeholder if fallback_payload is not None else svg_placeholder
        return fallback_reference, svg_placeholder, media_name or fallback_media_name

    def save(self):
        buffer = BytesIO()
        self._node.save(buffer)
        data = buffer.getvalue()
        if self._embedded_fonts:
            data = self._embedFonts(data)
        if self._inline_svg_pictures:
            data = self._embedPictures(data)
        self.output.write(data)
        if hasattr(self.output, "close"):
            self.output.close()

    def _embedFonts(self, data):
        with zipfile.ZipFile(BytesIO(data), "r") as zin:
            font_table = ET.fromstring(zin.read("word/fontTable.xml"))
            settings = ET.fromstring(zin.read("word/settings.xml"))
            content_types = ET.fromstring(zin.read("[Content_Types].xml"))
            try:
                rels = ET.fromstring(zin.read("word/_rels/fontTable.xml.rels"))
            except KeyError:
                rels = ET.Element(f"{{{_REL_NS}}}Relationships")
            replacements, font_parts = self._fontPackageParts(
                zin,
                font_table,
                settings,
                content_types,
                rels,
            )
            out = BytesIO()
            with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                skip = set(replacements) | set(font_parts)
                for item in zin.infolist():
                    if item.filename in skip:
                        continue
                    zout.writestr(item, zin.read(item.filename))
                for name, payload in replacements.items():
                    zout.writestr(name, payload)
                for name, payload in font_parts.items():
                    zout.writestr(name, payload)
            return out.getvalue()

    def _embedPictures(self, data):
        with zipfile.ZipFile(BytesIO(data), "r") as zin:
            content_types = ET.fromstring(zin.read("[Content_Types].xml"))
            document_xml = zin.read("word/document.xml").decode("utf-8")
            try:
                rels = ET.fromstring(zin.read("word/_rels/document.xml.rels"))
            except KeyError:
                rels = ET.Element(f"{{{_REL_NS}}}Relationships")
            replacements, media_parts = self._picturePackageParts(
                zin,
                document_xml,
                content_types,
                rels,
            )
            out = BytesIO()
            with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                skip = set(replacements) | set(media_parts)
                for item in zin.infolist():
                    if item.filename in skip:
                        continue
                    zout.writestr(item, zin.read(item.filename))
                for name, payload in replacements.items():
                    zout.writestr(name, payload)
                for name, payload in media_parts.items():
                    zout.writestr(name, payload)
            return out.getvalue()

    def _picturePackageParts(self, package, document_xml, content_types, rels):
        next_rid = self._nextRelationshipId(rels)
        existing_media = {
            name for name in package.namelist()
            if name.startswith("word/media/")
        }
        media_parts = {}
        relationship_ids = {}
        for picture in self._inline_svg_pictures.values():
            if picture.fallback_placeholder != picture.svg_placeholder:
                fallback_rid = f"rId{next_rid}"
                next_rid += 1
                fallback_media_name = self._uniqueMediaName(
                    picture.fallback_media_name,
                    existing_media | set(media_parts),
                )
                fallback_target = f"media/{fallback_media_name}"
                fallback_part_name = f"word/{fallback_target}"
                relationship_ids[picture.fallback_placeholder] = fallback_rid
                media_parts[fallback_part_name] = picture.fallback_payload
                self._appendImageRelationship(rels, fallback_rid, fallback_target)

            if picture.svg_placeholder is not None:
                svg_media_name = self._uniqueMediaName(
                    picture.media_name,
                    existing_media | set(media_parts),
                )
                svg_target = f"media/{svg_media_name}"
                svg_part_name = f"word/{svg_target}"
                svg_rid = f"rId{next_rid}"
                next_rid += 1
                relationship_ids[picture.svg_placeholder] = svg_rid
                media_parts[svg_part_name] = picture.payload
                self._appendImageRelationship(rels, svg_rid, svg_target)
        for placeholder, rid in sorted(relationship_ids.items(), key=lambda item: len(item[0]), reverse=True):
            document_xml = document_xml.replace(placeholder, rid)
        self._ensureContentType(content_types, "svg", _SVG_CONTENT_TYPE)
        self._ensureContentType(content_types, "png", "image/png")
        return {
            "word/document.xml": document_xml.encode("utf-8"),
            "word/_rels/document.xml.rels": _xml_bytes(rels),
            "[Content_Types].xml": _xml_bytes(content_types),
        }, media_parts

    def _fontPackageParts(self, package, font_table, settings, content_types, rels):
        existing_font_parts = {
            name for name in package.namelist()
            if name.startswith("word/fonts/font") and name.rsplit(".", 1)[-1] in {"odttf", "odttc"}
        }
        next_font = self._nextNumber(existing_font_parts, r"font(\d+)\.odtt[fc]$")
        next_rid = self._nextRelationshipId(rels)
        font_parts = {}
        extensions = set()
        for embedded in self._embedded_fonts.values():
            font_key = "{" + str(uuid.uuid4()).upper() + "}"
            payload, suffix = _docx_font_payload(embedded.path, embedded.font_number)
            part_name = f"word/fonts/font{next_font}{suffix}"
            next_font += 1
            target = f"fonts/{Path(part_name).name}"
            rid = f"rId{next_rid}"
            next_rid += 1
            font_parts[part_name] = _obfuscate_font(payload, font_key)
            extensions.add(suffix[1:])
            self._appendFontRelationship(rels, rid, target)
            self._appendFontTableEntry(
                font_table,
                embedded.family,
                embedded.face,
                rid,
                font_key,
            )
        for extension in extensions:
            self._ensureContentType(content_types, extension)
        self._ensureEmbedTrueTypeFonts(settings)
        self._preserveIgnorableNamespace(font_table, "w14", _W14_NS)
        return {
            "word/fontTable.xml": _xml_bytes(font_table),
            "word/_rels/fontTable.xml.rels": _xml_bytes(rels),
            "word/settings.xml": _xml_bytes(settings),
            "[Content_Types].xml": _xml_bytes(content_types),
        }, font_parts

    @staticmethod
    def _uniqueMediaName(name, used):
        base, suffix = os.path.splitext(name)
        candidate = name
        index = 2
        while f"word/media/{candidate}" in used:
            candidate = f"{base}-{index}{suffix}"
            index += 1
        return candidate

    @staticmethod
    def _nextNumber(names, pattern):
        values = []
        regex = re.compile(pattern)
        for name in names:
            match = regex.search(name)
            if match is not None:
                values.append(int(match.group(1)))
        return max(values, default=0) + 1

    @staticmethod
    def _nextRelationshipId(rels):
        ids = []
        for rel in rels.findall(f"{{{_REL_NS}}}Relationship"):
            rid = rel.get("Id")
            if rid and rid.startswith("rId") and rid[3:].isdigit():
                ids.append(int(rid[3:]))
        return max(ids, default=0) + 1

    @staticmethod
    def _appendFontRelationship(rels, rid, target):
        rel = ET.SubElement(rels, f"{{{_REL_NS}}}Relationship")
        rel.set("Id", rid)
        rel.set("Type", _FONT_REL_TYPE)
        rel.set("Target", target)

    @staticmethod
    def _appendImageRelationship(rels, rid, target):
        rel = ET.SubElement(rels, f"{{{_REL_NS}}}Relationship")
        rel.set("Id", rid)
        rel.set("Type", _IMAGE_REL_TYPE)
        rel.set("Target", target)

    @staticmethod
    def _appendFontTableEntry(font_table, family, face, rid, font_key):
        font = None
        for candidate in font_table.findall(f"{{{_W_NS}}}font"):
            if candidate.get(f"{{{_W_NS}}}name") == family:
                font = candidate
                break
        if font is None:
            font = ET.SubElement(font_table, f"{{{_W_NS}}}font")
            font.set(f"{{{_W_NS}}}name", family)
        embed_name = _WORD_FONT_FACE_ELEMENTS[face]
        for child in list(font):
            if child.tag == f"{{{_W_NS}}}{embed_name}":
                font.remove(child)
        embed = ET.SubElement(font, f"{{{_W_NS}}}{embed_name}")
        embed.set(f"{{{_R_NS}}}id", rid)
        embed.set(f"{{{_W_NS}}}fontKey", font_key)

    @staticmethod
    def _ensureContentType(content_types, extension, content_type=_OBFUSCATED_FONT_CONTENT_TYPE):
        for default in content_types.findall(f"{{{_CT_NS}}}Default"):
            if default.get("Extension") == extension:
                default.set("ContentType", content_type)
                return
        default = ET.SubElement(content_types, f"{{{_CT_NS}}}Default")
        default.set("Extension", extension)
        default.set("ContentType", content_type)

    @staticmethod
    def _ensureEmbedTrueTypeFonts(settings):
        tag = f"{{{_W_NS}}}embedTrueTypeFonts"
        if settings.find(tag) is None:
            settings.append(ET.Element(tag))

    @staticmethod
    def _preserveIgnorableNamespace(node, prefix, uri):
        ignorable = node.get(f"{{{_MC_NS}}}Ignorable")
        if ignorable and prefix in ignorable.split():
            node.set(f"xmlns:{prefix}", uri)


class DocxBackend(reflow.Reflow):
    """
    Very small proof-of-concept DOCX backend.

    Scope intentionally stays narrow:
    - page-wise reconstruction from shipped TeX pages
    - TeX controls both paragraph and line breaks
    - paragraphs, basic alignments, and inline/display math are supported
    - images, specials, and broader page semantics remain intentionally narrow

    The backend reconstructs TeX paragraphs from shipped line boxes, emits one
    Word paragraph per TeX paragraph, inserts explicit line breaks between TeX
    lines, and uses TeX's already-computed paragraph ownership and glue as DOCX
    paragraph spacing hints.
    """

    support_annotation = True
    supported_graphic_formats = ("svg",)

    def __init__(self, parser, output=None):
        super().__init__(parser, paginate=True)
        self.output = output
        self.file = None
        self.finished = False
        self._docx_next_drawing_id = 1
        self._docx_next_textbox_id = 1
        self.section = None
        self.docx_path = None
    
    def open(self):
        output = self.parser.jobname
        output = os.fspath(output)
        if output.startswith("./"):
            output = output[2:]
        if not output.endswith(".docx"):
            output += ".docx"
        if not self.parser.resolver.output_in_memory:
            self.docx_path = Path(self.parser.resolver._outputPath(output))
        return Document(self.parser.jobname, self.parser.resolver.openOut(output, "shipout/docx"))

    def _annotation_link(self, name=None, payload=None):
        anchor = None
        href = None
        if payload:
            info = _annotation_info(payload)
            kind = info.get("kind")
            if kind == "goto":
                anchor = info["destination"]
            elif kind == "gotor":
                href = info["file"]
                if info.get("destination"):
                    href += "#" + info["destination"]
            elif kind == "uri":
                href = info["url"]
        if anchor is None and href is None and name is not None:
            if name.startswith("#"):
                anchor = name[1:]
            elif "#" in name or ":" in name:
                href = name
            else:
                anchor = name.lstrip("@")
        reopen_name = href or (None if anchor is None else "#" + anchor) or name or ""
        return reopen_name, anchor, href

    def newAnnotationBuilder(self, name=None, payload=None):
        reopen_name, anchor, href = self._annotation_link(name=name, payload=payload)
        return AnnotationBuilder(self, self.builder, reopen_name, anchor=anchor, href=href)

    def newFixedAnnotation(self, name, w, h):
        return None

    def _preserve_alignment_edge_glue(self, line_box):
        container = getattr(self.builder, "container", None)
        story = getattr(container, "story", None)
        return isinstance(story, Story)

    def setTarget(self, name):
        if self.document is None or self.builder is None:
            return
        container = getattr(self.builder, "container", None)
        paragraph = getattr(container, "_node", None)
        if isinstance(container, AnnotationLine):
            paragraph = container.parent._node
        if paragraph is None or not hasattr(paragraph, "_p"):
            return
        bookmark_id = str(self.document.nextBookmarkId())
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), bookmark_id)
        start.set(qn("w:name"), name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), bookmark_id)
        paragraph._p.append(start)
        paragraph._p.append(end)

    def _hbox_extent(self, box):
        x = Dimen()
        left = None
        right = None
        glue_state = self._glue_state(box)
        for node in getattr(box, "list", ()):
            node_type = getattr(node, "node_type", None)
            if node_type == nd.NODE_TYPE.GLUE:
                if glue_state is None:
                    x += node.glue.dimen
                else:
                    x += Dimen(integer=self._glue_amount(node, None, glue_state))
                continue
            if node_type == nd.NODE_TYPE.KERN:
                x += node.kern
                continue
            if node_type == nd.NODE_TYPE.PENALTY:
                continue
            if node_type == nd.NODE_TYPE.MATH:
                x += getattr(node, "kern", Dimen())
                continue
            node_left, node_right = self._node_extent(node)
            if left is None or x + node_left < left:
                left = x + node_left
            if right is None or x + node_right > right:
                right = x + node_right
            width = getattr(node, "width", None)
            if width is not None:
                x += width
        if left is None:
            return Dimen(), Dimen()
        return left, right

    def _node_extent(self, node):
        if getattr(node, "node_type", None) == nd.NODE_TYPE.HLIST:
            return self._hbox_extent(node)
        width = getattr(node, "width", None)
        if width is None:
            return Dimen(), Dimen()
        return Dimen(), Dimen(width)

    def _visible_hbox_content(self, box):
        content = []
        for node in getattr(box, "list", ()):
            node_type = getattr(node, "node_type", None)
            if node_type in (nd.NODE_TYPE.GLUE, nd.NODE_TYPE.KERN, nd.NODE_TYPE.PENALTY):
                continue
            if node_type == nd.NODE_TYPE.HLIST:
                left, right = self._hbox_extent(node)
                if int(getattr(node, "width", Dimen())) == 0 and int(right - left) > 0:
                    content.extend(self._visible_hbox_content(node))
                else:
                    content.append(node)
                continue
            content.append(node)
        return content

    def _alignment_cell_docx_box(self, cell):
        width = self._alignment_cell_width(cell)
        if width is None or int(width) != 0:
            return cell, width, Dimen()
        left, right = self._hbox_extent(cell)
        visible = right - left
        if int(visible) <= 0:
            return cell, width, Dimen()
        render = cell
        if int(left) < 0 and hasattr(cell, "copy"):
            content = self._visible_hbox_content(cell)
            render = cell.copy(content=content) if content else cell.copy()
            render.width = visible
            render.to = visible
            render.spread = Dimen()
            render.glue_ratio = bx.GlueRatio(0, 0, 1)
        return render, visible, -left if int(left) < 0 else Dimen()

    def typesetHeaderRegion(self, items):
        self._typesetVModePageRegion(self.document.header, items)

    def typesetFooterRegion(self, items):
        self._typesetVModePageRegion(self.document.footer, items)

    def _typesetVModePageRegion(self, story, items):
        region_items = [
            item for item in items
            if getattr(item.node, "node_type", None) in (nd.NODE_TYPE.HLIST, nd.NODE_TYPE.VLIST, nd.NODE_TYPE.WHATSIT)
        ]
        if not region_items:
            return

        box_items = [
            item for item in region_items
            if getattr(item.node, "node_type", None) in (nd.NODE_TYPE.HLIST, nd.NODE_TYPE.VLIST)
        ]
        render_box_items = [
            item for item in box_items
            if self._region_node_has_layout(item.node)
        ]
        if render_box_items and hasattr(story, "clear"):
            story.clear()
        base_x = min((item.x for item in box_items), default=Dimen())

        with reflow.Builder(self, story):
            for item in region_items:
                node = item.node
                node_type = getattr(node, "node_type", None)
                if node_type == nd.NODE_TYPE.WHATSIT:
                    node.output(self.parser, self)
                    continue
                if not self._region_node_has_layout(node):
                    self.scanWhatsits([node])
                    continue
                spacing = item.y
                if int(spacing) < 0:
                    spacing = Dimen()
                xspacing = item.x - base_x
                if node_type == nd.NODE_TYPE.VLIST:
                    self.typesetVBox(node, xspacing=xspacing, yspacing=spacing)
                elif node_type == nd.NODE_TYPE.HLIST:
                    self.typesetHBox(node, xspacing=xspacing, yspacing=spacing)

    def typesetInlineVBox(self, box: bx.Box):
        if int(box.width) == 0:
            return None
        block = super().typesetInlineVBox(box)
        if isinstance(block, TextBoxStory):
            block.finalizeContent()
        return block

    def typesetInlineMath(self, node: mmode.InlineMathNode, box: bx.Box, piece: int):
        self._require_builder("typesetInlineMath", "newInlineMath")
        return self.builder.newInlineMath(self, node, box, piece)

    def typesetGraphicAsset(self, asset, request):
        if asset.format != "svg" or self.builder is None:
            return
        new_text_run = getattr(self.builder, "newTextRun", None)
        if new_text_run is not None:
            baseline_from_bottom = self._current_run_baseline_from_bottom
            if baseline_from_bottom is None:
                baseline_from_bottom = getattr(self.builder, "baseline_from_bottom", Dimen())
            run = new_text_run(
                text=None,
                font=self.parser.parameters["currentfont"],
                color=self.color,
                baseline_from_bottom=baseline_from_bottom,
            )
            with reflow.Builder(self, run):
                graphic = self.builder.newInlineGraphic(self, asset, request)
            if graphic is not None:
                return request.width or asset.width

    @staticmethod
    def graphicSvgPayload(asset, width=None, height=None):
        if asset.data is not None:
            if isinstance(asset.data, str):
                payload = asset.data.encode("utf-8")
            else:
                payload = asset.data
            if width is not None and height is not None:
                payload = _retarget_svg_size(payload, Dimen(width), Dimen(height))
            return payload
        if asset.path is None:
            return None
        try:
            payload = Path(asset.path).read_bytes()
        except OSError:
            return None
        if width is not None and height is not None:
            payload = _retarget_svg_size(payload, Dimen(width), Dimen(height))
        return payload

    def typesetTrailingVListSpacing(self, spacing: Dimen, top_level: bool=False):
        if not top_level or int(spacing) >= 0:
            return
        apply_spacing = getattr(self.builder, "applyTrailingSpacing", None)
        if apply_spacing is not None:
            apply_spacing(spacing)

    def typesetDisplayMath(self, node: mmode.DisplayMathNode, collection, yspacing: Dimen=Dimen(), glue_state=None):
        self._require_builder("typesetDisplayMath", "newParagraph")
        display_boxes, trailing_spacing = self._display_math_boxes(collection, yspacing, glue_state=glue_state)
        for box, spacing_before in display_boxes:
            self._typesetDisplayMathBox(node, box, spacing_before)
        return trailing_spacing

    def _display_math_boxes(self, collection, yspacing, glue_state=None):
        boxes = []
        pending_spacing = Dimen(yspacing)
        for n in collection:
            node_type = getattr(n, "node_type", None)
            if node_type == nd.NODE_TYPE.HLIST:
                boxes.append((n, pending_spacing))
                pending_spacing = Dimen()
                continue
            if node_type == nd.NODE_TYPE.GLUE:
                glue = getattr(n, "glue", None)
                if glue is not None:
                    if glue_state is None:
                        pending_spacing += glue.dimen
                    else:
                        pending_spacing += Dimen(integer=self._glue_amount(n, None, glue_state))
                continue
            if node_type == nd.NODE_TYPE.KERN:
                pending_spacing += n.kern
                continue
            if node_type == nd.NODE_TYPE.WHATSIT:
                n.output(self.parser, self)
        return boxes, pending_spacing

    def _typesetDisplayMathBox(self, node, box, spacing_before):
        picture_box = DisplayMathPictureBox(box)
        para = self.builder.newParagraph(spacing_before=spacing_before, justify="left")
        with reflow.ParagraphBuilder(self, para):
            line_spec = reflow.LineSpec(self, picture_box, spacing_before=Dimen())
            line = para.newLine(line_spec)
            with reflow.LineBuilder(self, line):
                text_run = self.builder.newTextRun(
                    text=None,
                    font=self.parser.parameters["currentfont"],
                    color=self.color,
                    baseline_from_bottom=picture_box.depth,
                )
                with reflow.Builder(self, text_run):
                    self.typesetInlineMath(node, picture_box, 1)

    def inlineMathSvg(self, box: bx.Box):
        with tempfile.TemporaryDirectory() as tmpdir:
            prefix = Path(tmpdir) / "inline-math"
            hoffset = self.parser.layout["hoffset"]
            voffset = self.parser.layout["voffset"]
            try:
                self.parser.layout["hoffset"] = Dimen()
                self.parser.layout["voffset"] = Dimen()
                backend = svg.SVGShipoutBackend(self.parser, os.fspath(prefix))
                backend.shipout(box)
            finally:
                self.parser.layout["hoffset"] = hoffset
                self.parser.layout["voffset"] = voffset
            return Path(f"{prefix}-1.svg").read_bytes()

    def define_font(self, font):
        if font is not None:
            _require_opentype_font_backend(font)
        if self.document is None:
            return None
        reference = self.document.defineFont(font)
        return None if reference is None else reference.family


def init(parser):
    parser.shipout = DocxBackend(parser)
    parser.font_size_in_bp = True


mod = Module(
    "docx",
    init=init,
    attributes={},
)
