import io
import math
import re
import uuid
import zipfile
import xml.etree.ElementTree as ET

import pytest
from docx import Document as WordDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fontTools.fontBuilder import FontBuilder
from fontTools.pens.boundsPen import BoundsPen
from fontTools.pens.t2CharStringPen import T2CharStringPen
from fontTools.pens.ttGlyphPen import TTGlyphPen
from fontTools.ttLib import TTFont

from pytex import align
from pytex import box as bx
from pytex import docx
from pytex import graphics
from pytex import font as txfont
from pytex import mmode
from pytex import node as nd
from pytex import opentype
from pytex import paragraph as pg
from pytex import reflow
from pytex.dimen import Dimen
from pytex.font_backend import GlyphInfo
from pytex.glue import Glue, Stretchness
from pytex.parser import Parser
from pytex.token import CATCODE


@pytest.fixture()
def parser(tmp_path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    p = Parser()
    p.resolver.output_in_memory = True
    for ch, cat in [
        ("{", CATCODE.BEGIN_GROUP),
        ("}", CATCODE.END_GROUP),
        ("$", CATCODE.MATH_SHIFT),
        ("&", CATCODE.ALIGNMENT_TAB),
        ("#", CATCODE.PARAMETER),
        ("^", CATCODE.SUPERSCRIPT),
        ("_", CATCODE.SUBSCRIPT),
    ]:
        p.catcode[ord(ch)] = cat
    p.layout["hsize"] = Dimen(200)
    p.layout["vsize"] = Dimen(300)
    yield p
    p.close()


class _FakeBackend:
    def __init__(self, name="Fake Roman", kind="opentype", path=None, font_number=0):
        self.name = name
        self.kind = kind
        self.path = path
        self.font_number = font_number
        self.fontdimen = [0.0, 0.5, 0.25, 0.15, 0.7, 1.0, 0.0]
        self.units_per_em = 1000
        self.font = {
            "hhea": type(
                "FakeHhea",
                (),
                {
                    "ascent": 1000,
                    "descent": 0,
                    "lineGap": 0,
                },
            )()
        }

    def glyphInfo(self, char):
        return GlyphInfo(char=char, width=0.5, height=0.7, depth=0.2, italic=0)

    def fallbackGlyphInfo(self, char):
        return self.glyphInfo(char)

    def hasChar(self, char):
        return True

    def _spaceWidth(self):
        return 0.5


class _FakeWordBaselineBackend(_FakeBackend):
    def __init__(self, baseline, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.baseline = baseline

    def lineBaselineFromBottom(self, font_size, line_height, round_total=None):
        return self.baseline


class _FakeFont(txfont.Font):
    def __init__(self, name="Fake Roman", size=10, kind="opentype", path=None, font_number=0):
        super().__init__(_FakeBackend(name, kind=kind, path=path, font_number=font_number), Dimen(size))


class _FakeWordBaselineFont(txfont.Font):
    def __init__(self, baseline, name="Fake Roman", size=10, kind="opentype", path=None, font_number=0):
        super().__init__(
            _FakeWordBaselineBackend(
                baseline,
                name=name,
                kind=kind,
                path=path,
                font_number=font_number,
            ),
            Dimen(size),
        )


class _FakeHBox:
    node_type = nd.NODE_TYPE.HLIST

    def __init__(
        self,
        items,
        source=None,
        width=80,
        height=7,
        depth=2,
        rightmost_value=None,
        natural=None,
        glue_ratio=None,
    ):
        self.list = items
        self.source = source
        self.width = Dimen(width)
        self.height = Dimen(height)
        self.depth = Dimen(depth)
        self.to = Dimen(width)
        self.spread = Dimen()
        self.natural = Glue() if natural is None else natural
        self.glue_ratio = Dimen() if glue_ratio is None else glue_ratio
        self.shifted = Dimen()
        self._rightmost_value = self.width if rightmost_value is None else Dimen(rightmost_value)

    def rightmost(self):
        return self._rightmost_value


class _FakeVBox:
    node_type = nd.NODE_TYPE.VLIST

    def __init__(self, items, width=200, height=120, depth=0):
        self.list = items
        self.source = None
        self.width = Dimen(width)
        self.height = Dimen(height)
        self.depth = Dimen(depth)
        self.to = Dimen(height)
        self.spread = Dimen()
        self.natural = Glue()
        self.glue_ratio = Dimen()
        self.shifted = Dimen()


def _char_nodes(text, font):
    return [nd.CharNode(ch, font) for ch in text]


def _line_box(text, source, font=None, width=80):
    font = _FakeFont() if font is None else font
    nodes = []
    for ch in text:
        if ch == " ":
            nodes.append(nd.Glue(Glue(font.param[1]), None))
        else:
            nodes.append(nd.CharNode(ch, font))
    return _FakeHBox(nodes, source=source, width=width)


def _page_box(items):
    return _FakeVBox([nd.Glue(Glue(Dimen(10)), "\\topskip"), *items])


def _docx_bytes(parser, backend):
    backend.close()
    return parser.resolver.in_memory_files["texput.docx"].content


def _word_document(parser, backend):
    return WordDocument(io.BytesIO(_docx_bytes(parser, backend)))


def _document_xml(data):
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        return zf.read("word/document.xml").decode("utf-8")


def _drawing_runs(xml):
    return [
        match.group(0)
        for match in re.finditer(
            r"<w:r>(?:(?!</w:r>).)*<w:drawing(?:(?!</w:r>).)*</w:r>",
            xml,
            flags=re.S,
        )
    ]


def _document_root(data):
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        return ET.fromstring(zf.read("word/document.xml"))


def _document_relationships_root(data):
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        return ET.fromstring(zf.read("word/_rels/document.xml.rels"))


def _install_font(parser, font=None):
    font = _FakeFont() if font is None else font
    parser.parameters["currentfont"] = font
    return font


def _build_test_word_font(path, family, subfamily, *, bold=False, italic=False):
    builder = FontBuilder(1000, isTTF=True)
    glyph_order = [".notdef", "A", "B", "I", "R", "T"]
    builder.setupGlyphOrder(glyph_order)
    builder.setupCharacterMap({ord(char): char for char in glyph_order[1:]})
    glyphs = {}
    for glyph_name in glyph_order:
        pen = TTGlyphPen(None)
        if glyph_name != ".notdef":
            pen.moveTo((80, 0))
            pen.lineTo((300, 700))
            pen.lineTo((520, 0))
            pen.closePath()
        glyphs[glyph_name] = pen.glyph()
    builder.setupGlyf(glyphs)
    builder.setupHorizontalMetrics({glyph_name: (600, 0) for glyph_name in glyph_order})
    builder.setupHorizontalHeader(ascent=800, descent=-200)
    builder.setupNameTable(
        {
            "familyName": family,
            "styleName": subfamily,
            "uniqueFontIdentifier": f"{family} {subfamily}",
            "fullName": f"{family} {subfamily}",
            "psName": re.sub(r"[^A-Za-z0-9-]", "", f"{family}-{subfamily}"),
        }
    )
    fs_selection = (1 << 5 if bold else 0) | (1 << 0 if italic else 0)
    if not bold and not italic:
        fs_selection |= 1 << 6
    builder.setupOS2(
        sTypoAscender=800,
        sTypoDescender=-200,
        usWinAscent=800,
        usWinDescent=200,
        usWeightClass=700 if bold else 400,
        fsSelection=fs_selection,
    )
    builder.setupPost(italicAngle=-12 if italic else 0)
    builder.setupMaxp()
    builder.font["head"].macStyle = (1 if bold else 0) | (2 if italic else 0)
    builder.save(path)


def _build_test_cff_word_font(path, family, subfamily):
    builder = FontBuilder(1000, isTTF=False)
    glyph_order = [".notdef", "A"]
    builder.setupGlyphOrder(glyph_order)
    builder.setupCharacterMap({ord("A"): "A"})
    builder.setupHorizontalMetrics({glyph_name: (600, 0) for glyph_name in glyph_order})
    builder.setupHorizontalHeader(ascent=800, descent=-200)
    builder.setupNameTable(
        {
            "familyName": family,
            "styleName": subfamily,
            "uniqueFontIdentifier": f"{family} {subfamily}",
            "fullName": f"{family} {subfamily}",
            "psName": re.sub(r"[^A-Za-z0-9-]", "", f"{family}-{subfamily}"),
        }
    )
    builder.setupOS2(
        sTypoAscender=800,
        sTypoDescender=-200,
        usWinAscent=800,
        usWinDescent=200,
    )
    builder.setupPost()
    char_strings = {}
    for glyph_name in glyph_order:
        pen = T2CharStringPen(600, None)
        if glyph_name != ".notdef":
            pen.moveTo((80, 0))
            pen.lineTo((300, 700))
            pen.lineTo((520, 0))
            pen.closePath()
        char_strings[glyph_name] = pen.getCharString()
    builder.setupCFF(
        re.sub(r"[^A-Za-z0-9-]", "", f"{family}-{subfamily}"),
        {"FullName": f"{family} {subfamily}", "FamilyName": family, "Weight": subfamily},
        char_strings,
        {},
    )
    builder.save(path)


def _font_from_word_metadata(path, tex_name, size=10):
    font = _FakeFont(tex_name, size=size, kind="opentype", path=str(path))
    font.backend.font = TTFont(path)
    font.backend.units_per_em = font.backend.font["head"].unitsPerEm
    return font


def _text_position_to_tex_baseline(run_baseline, backend_baseline):
    return docx.half_pt(Dimen(run_baseline) - Dimen(backend_baseline))


def _box_position_to_tex_baseline(line_baseline, run_baseline, backend_baseline):
    return docx.half_pt(Dimen(line_baseline) - Dimen(backend_baseline) - Dimen(run_baseline))


def _math_atom(char, fam=0, atom_type=mmode.ATOM_TYPE.ORD):
    atom = mmode.Atom(atom_type)
    atom.nucleus = mmode.MathSymbol((atom_type.value << 12) | (fam << 8) | ord(char), -1)
    return atom


def test_docx_twips_truncate_to_word_unit():
    assert docx._twips(docx._tex_points(0.099)) == 1
    assert docx._twips(docx._tex_points(-0.099)) == -1
    assert docx._twips(docx._tex_points(0.049)) == 0
    assert docx._twips(docx._tex_points(-0.049)) == 0
    assert docx.half_pt(docx._tex_points(0.26)) == "1"
    assert docx.half_pt(docx._tex_points(-0.26)) == "-1"


def test_docx_space_twips_floor_to_preserve_line_fit():
    assert docx._space_twips(docx._tex_points(0.099)) == 1
    assert docx._space_twips(docx._tex_points(-0.099)) == -2
    assert docx._space_twips(docx._tex_points(0.049)) == 0
    assert docx._space_twips(docx._tex_points(-0.049)) == -1


def _alignment_cell(text, font, width=20):
    box = _FakeHBox(_char_nodes(text, font), width=width, rightmost_value=width)
    box.span = 1
    return box


def _alignment_row_box(owner, row, width=None, height=8, depth=2):
    nodes = []
    natural_width = Dimen()
    for index, cell in enumerate(row.cells):
        if index < len(owner.tabskips):
            glue_node = nd.Glue(owner.tabskips[index], None)
            nodes.append(glue_node)
            natural_width += owner.tabskips[index].dimen
        nodes.append(cell)
        natural_width += cell.width
    if len(owner.tabskips) > len(row.cells):
        glue_node = nd.Glue(owner.tabskips[len(row.cells)], None)
        nodes.append(glue_node)
        natural_width += owner.tabskips[len(row.cells)].dimen
    return _FakeHBox(
        nodes,
        source=owner,
        width=natural_width if width is None else width,
        height=height,
        depth=depth,
    )


def test_docx_init_selects_reflow_backend_and_bp_font_sizes(parser):
    docx.init(parser)

    assert isinstance(parser.shipout, docx.DocxBackend)
    assert parser.font_size_in_bp is True


def test_docx_define_font_requires_opentype_shape(parser):
    backend = docx.DocxBackend(parser)
    font = _FakeFont(kind="tfm")

    with pytest.raises(AssertionError, match="OpenType-shaped"):
        backend.define_font(font)


def test_docx_line_baseline_uses_hhea_padding_scaled_to_exact_line_height(parser):
    font = _FakeFont()
    hhea = font.backend.font["hhea"]
    hhea.ascent = 800
    hhea.descent = -200
    hhea.lineGap = 100
    line_box = _FakeHBox([], width=20, height=18, depth=6)
    line_spec = type(
        "FakeLineSpec",
        (),
        {
            "line_height": line_box.height + line_box.depth,
            "spacing_before": Dimen(),
            "default_font": font,
            "line_box": line_box,
            "baseline_from_bottom": line_box.depth,
        },
    )()
    line = docx.Line(WordDocument().add_paragraph(), 1, line_spec)

    font_size = round(docx._docx_points(font.at) * 2) / 2
    padding = round(0.15 * (hhea.ascent - hhea.descent))
    total_units = hhea.ascent - hhea.descent + hhea.lineGap + 2 * padding
    total_size = math.ceil(total_units / font.backend.units_per_em * font_size * 2 - 1e-9) / 2
    expected = docx._docx_points(line_spec.line_height) * (
        (-hhea.descent + padding) / font.backend.units_per_em * font_size
    ) / total_size

    assert line.backendBaselineForFont(font) == pytest.approx(expected)


def test_docx_document_interface_uses_pagespec_sections(parser):
    backend = docx.DocxBackend(parser)
    document = backend.open()

    section = document.newPage(
        reflow.PageSpec(
            width=Dimen(100),
            height=Dimen(200),
            margin_left=Dimen(10),
            margin_top=Dimen(20),
            margin_right=Dimen(10),
            margin_bottom=Dimen(20),
            header_distance=Dimen(8),
            footer_distance=Dimen(12),
        )
    )

    assert isinstance(document, reflow.Document)
    assert isinstance(section.body, reflow.Element)
    assert document.body is section.body
    word_section = document._node.sections[0]
    assert int(word_section.page_width) == int(docx._length(Dimen(100)))
    assert int(word_section.page_height) == int(docx._length(Dimen(200)))
    assert int(word_section.left_margin) == int(docx._length(Dimen(10)))
    assert int(word_section.top_margin) == int(docx._length(Dimen(20)))
    assert int(word_section.right_margin) == int(docx._length(Dimen(10)))
    assert int(word_section.bottom_margin) == int(docx._length(Dimen(20)))
    assert int(word_section.header_distance) == int(docx._length(Dimen(8)))
    assert int(word_section.footer_distance) == int(docx._length(Dimen(12)))


def test_docx_section_break_uses_minimized_empty_paragraph(parser):
    backend = docx.DocxBackend(parser)
    document = backend.open()
    spec = reflow.PageSpec(
        width=Dimen(100),
        height=Dimen(200),
        margin_left=Dimen(10),
        margin_top=Dimen(20),
        margin_right=Dimen(10),
        margin_bottom=Dimen(20),
    )
    document.newPage(spec)
    document._node.add_paragraph("Visible")

    document.newPage(spec)

    buffer = io.BytesIO()
    document._node.save(buffer)
    with zipfile.ZipFile(io.BytesIO(buffer.getvalue())) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    body = root.find(f"{{{docx._W_NS}}}body")
    paragraphs = [child for child in body if child.tag == f"{{{docx._W_NS}}}p"]
    visible_index = next(
        index for index, para in enumerate(paragraphs)
        if "".join(t.text or "" for t in para.findall(f".//{{{docx._W_NS}}}t")) == "Visible"
    )
    visible = paragraphs[visible_index]
    section_break = paragraphs[visible_index + 1]

    assert visible.find(f"{{{docx._W_NS}}}pPr/{{{docx._W_NS}}}sectPr") is None
    assert section_break.find(f"{{{docx._W_NS}}}pPr/{{{docx._W_NS}}}sectPr") is not None
    spacing = section_break.find(f"{{{docx._W_NS}}}pPr/{{{docx._W_NS}}}spacing")
    assert spacing is not None
    assert spacing.get(f"{{{docx._W_NS}}}before") == "0"
    assert spacing.get(f"{{{docx._W_NS}}}after") == "0"
    assert spacing.get(f"{{{docx._W_NS}}}line") == "1"
    assert spacing.get(f"{{{docx._W_NS}}}lineRule") == "exact"


def test_docx_shipout_embeds_filesystem_opentype_fonts(parser, tmp_path):
    font_bytes = bytes(range(64))
    font_path = tmp_path / "DemoFont.otf"
    font_path.write_bytes(font_bytes)
    font = _FakeFont("Demo Embedded", kind="opentype", path=str(font_path))
    _install_font(parser, font)
    owner = pg.Paragraph(parser, indent=False)
    backend = docx.DocxBackend(parser)
    parser.shipout = backend

    backend.shipout(_page_box([_line_box("Hi", owner, font)]))
    data = _docx_bytes(parser, backend)

    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        names = set(zf.namelist())
        assert "word/fonts/font1.odttf" in names
        font_xml = zf.read("word/fontTable.xml").decode("utf-8")
        rels_xml = zf.read("word/_rels/fontTable.xml.rels").decode("utf-8")
        content_types = zf.read("[Content_Types].xml").decode("utf-8")
        settings_xml = zf.read("word/settings.xml").decode("utf-8")
        obfuscated = zf.read("word/fonts/font1.odttf")

    assert 'w:name="Demo Embedded"' in font_xml
    assert "<w:embedRegular" in font_xml
    font_key = re.search(r'w:fontKey="([^"]+)"', font_xml).group(1)
    assert font_key.startswith("{") and font_key.endswith("}")
    assert 'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font"' in rels_xml
    assert 'Target="fonts/font1.odttf"' in rels_xml
    assert 'Extension="odttf"' in content_types
    assert "application/vnd.openxmlformats-officedocument.obfuscatedFont" in content_types
    assert "<w:embedTrueTypeFonts" in settings_xml
    key = docx._font_key_bytes(font_key)
    restored = bytearray(obfuscated)
    for index in range(32):
        restored[index] ^= key[index % 16]
    assert bytes(restored) == font_bytes


def test_docx_converts_embedded_cff_font_to_truetype(parser, tmp_path):
    font_path = tmp_path / "DemoCFF.otf"
    _build_test_cff_word_font(font_path, "Demo CFF", "Regular")
    parser.supported_font_classes = None
    source_backend = parser.loadFontBackend(str(font_path))
    assert isinstance(source_backend, opentype.CFFBackend)

    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    assert parser.supported_font_classes == (opentype.TrueTypeBackend,)
    converted_backend = parser.loadFontBackend(str(font_path))
    assert isinstance(converted_backend, opentype.TrueTypeBackend)
    assert "glyf" in converted_backend.font
    assert "loca" in converted_backend.font
    assert "CFF " not in converted_backend.font

    font = txfont.Font(converted_backend, Dimen(10))
    _install_font(parser, font)
    owner = pg.Paragraph(parser, indent=False)

    info = font.backend.glyphInfo("A")
    glyph = font.backend.font.getGlyphSet()["A"]
    bounds_pen = BoundsPen(font.backend.font.getGlyphSet())
    glyph.draw(bounds_pen)
    assert info.height == bounds_pen.bounds[3] / font.backend.units_per_em
    assert info.depth == -bounds_pen.bounds[1] / font.backend.units_per_em

    backend.shipout(_page_box([_line_box("A", owner, font)]))
    data = _docx_bytes(parser, backend)

    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        font_xml = zf.read("word/fontTable.xml").decode("utf-8")
        obfuscated = zf.read("word/fonts/font1.odttf")

    font_key = re.search(r'w:fontKey="([^"]+)"', font_xml).group(1)
    restored = bytearray(obfuscated)
    key = docx._font_key_bytes(font_key)
    for index in range(min(32, len(restored))):
        restored[index] ^= key[index % 16]
    assert bytes(restored) == converted_backend.fontData()
    converted = TTFont(io.BytesIO(restored))
    assert converted.sfntVersion == "\x00\x01\x00\x00"
    assert "glyf" in converted
    assert "loca" in converted
    assert "CFF " not in converted
    assert converted.getBestCmap()[ord("A")] == "A"
    converted.close()


def test_docx_keeps_truetype_font_payload_unchanged(parser, tmp_path):
    font_path = tmp_path / "DemoTrueType.ttf"
    _build_test_word_font(font_path, "Demo TrueType", "Regular")

    backend = parser.loadFontBackend(str(font_path))

    assert isinstance(backend, opentype.TrueTypeBackend)
    assert backend.fontData() == font_path.read_bytes()


def test_docx_groups_embedded_font_faces_by_word_family(parser, tmp_path):
    specs = {
        "regular": ("LM Roman 12", "Regular", False, False),
        "bold": ("LM Roman 12", "Bold", True, False),
        "italic": ("LM Roman 12", "Italic", False, True),
        "boldItalic": ("LM Roman 12", "Bold Italic", True, True),
        "title": ("LM Roman 17", "Regular", False, False),
    }
    fonts = {}
    for face, (family, subfamily, bold, italic) in specs.items():
        path = tmp_path / f"tex-{face}.ttf"
        _build_test_word_font(path, family, subfamily, bold=bold, italic=italic)
        fonts[face] = _font_from_word_metadata(path, f"tex-{face}", size=17 if face == "title" else 12)

    _install_font(parser, fonts["regular"])
    owner = pg.Paragraph(parser, indent=False)
    line = _FakeHBox(
        [
            nd.CharNode("R", fonts["regular"]),
            nd.CharNode("B", fonts["bold"]),
            nd.CharNode("I", fonts["italic"]),
            nd.CharNode("A", fonts["boldItalic"]),
            nd.CharNode("T", fonts["title"]),
        ],
        source=owner,
    )
    backend = docx.DocxBackend(parser)
    parser.shipout = backend

    backend.shipout(_page_box([line]))
    data = _docx_bytes(parser, backend)

    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        names = set(zf.namelist())
        document_root = ET.fromstring(zf.read("word/document.xml"))
        font_table = ET.fromstring(zf.read("word/fontTable.xml"))
        relationships = ET.fromstring(zf.read("word/_rels/fontTable.xml.rels"))

        family_entries = {
            entry.get(f"{{{docx._W_NS}}}name"): entry
            for entry in font_table.findall(f"{{{docx._W_NS}}}font")
        }
        roman12 = family_entries["LM Roman 12"]
        roman17 = family_entries["LM Roman 17"]
        assert "tex-regular" not in family_entries
        assert [
            child.tag.rsplit("}", 1)[-1]
            for child in roman12
            if child.tag.rsplit("}", 1)[-1].startswith("embed")
        ] == ["embedRegular", "embedBold", "embedItalic", "embedBoldItalic"]
        assert roman17.find(f"{{{docx._W_NS}}}embedRegular") is not None

        rel_targets = {
            rel.get("Id"): rel.get("Target")
            for rel in relationships.findall(f"{{{docx._REL_NS}}}Relationship")
        }
        for entry in (roman12, roman17):
            for embed in [child for child in entry if child.tag.rsplit("}", 1)[-1].startswith("embed")]:
                rid = embed.get(f"{{{docx._R_NS}}}id")
                target = rel_targets[rid]
                part_name = f"word/{target}"
                assert part_name in names
                font_key = embed.get(f"{{{docx._W_NS}}}fontKey")
                key = uuid.UUID(font_key.strip("{}")).bytes_le[::-1]
                restored = bytearray(zf.read(part_name))
                for index in range(min(32, len(restored))):
                    restored[index] ^= key[index % 16]
                parsed = TTFont(io.BytesIO(restored))
                assert parsed["name"].getName(1, 3, 1).toUnicode() == entry.get(f"{{{docx._W_NS}}}name")

    runs = {}
    for run in document_root.findall(f".//{{{docx._W_NS}}}r"):
        text = "".join(node.text or "" for node in run.findall(f"{{{docx._W_NS}}}t"))
        if text in "RBIAT":
            runs[text] = run
    assert set(runs) == set("RBIAT")

    for text in "RBIA":
        rfonts = runs[text].find(f"{{{docx._W_NS}}}rPr/{{{docx._W_NS}}}rFonts")
        assert rfonts is not None
        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
            assert rfonts.get(f"{{{docx._W_NS}}}{attr}") == "LM Roman 12"
    title_fonts = runs["T"].find(f"{{{docx._W_NS}}}rPr/{{{docx._W_NS}}}rFonts")
    assert title_fonts.get(f"{{{docx._W_NS}}}ascii") == "LM Roman 17"

    def has_style(text, tag):
        return runs[text].find(f"{{{docx._W_NS}}}rPr/{{{docx._W_NS}}}{tag}") is not None

    assert not has_style("R", "b") and not has_style("R", "i")
    assert has_style("B", "b") and not has_style("B", "i")
    assert not has_style("I", "b") and has_style("I", "i")
    assert has_style("A", "b") and has_style("A", "i")


def test_docx_shipout_writes_one_word_paragraph_per_tex_line(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    para1 = pg.Paragraph(parser, indent=False)
    para2 = pg.Paragraph(parser, indent=False)
    empty_source_glue = nd.Glue(Glue(Dimen()), None)
    empty_source_glue.source = pg.Paragraph(parser, indent=False)

    backend.shipout(
        _page_box(
            [
                _line_box("Hello world", para1, font),
                nd.Penalty(0),
                nd.Glue(Glue(Dimen(3)), "\\baselineskip"),
                empty_source_glue,
                _line_box("Again soon", para1, font),
                nd.Glue(Glue(Dimen(8)), "\\parskip"),
                _line_box("Second paragraph", para2, font),
            ]
        )
    )

    document = _word_document(parser, backend)
    assert [p.text.replace("\u00a0", " ") for p in document.paragraphs] == [
        "Hello world",
        "Again soon",
        "Second paragraph",
    ]
    assert int(document.paragraphs[1].paragraph_format.space_before) == int(docx._length(Dimen(3)))
    assert int(document.paragraphs[2].paragraph_format.space_before) == int(docx._length(Dimen(8)))


def test_docx_empty_hbox_width_becomes_nonbreaking_spacing(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    para = pg.Paragraph(parser, indent=False)
    empty = _FakeHBox([], width=12, height=0, depth=0, rightmost_value=12)
    line = _FakeHBox([nd.CharNode("A", font), empty, nd.CharNode("B", font)], para)

    backend.shipout(_page_box([line]))

    document = _word_document(parser, backend)
    assert document.paragraphs[0].text == "A\u00a0B"


def test_docx_paginated_leading_glue_becomes_paragraph_indent(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    para = pg.Paragraph(parser, indent=False)
    line = _FakeHBox(
        [
            nd.Glue(Glue(Dimen(18)), "\\parindent"),
            nd.CharNode("A", font),
        ],
        para,
    )

    backend.shipout(_page_box([line]))

    document = _word_document(parser, backend)
    assert document.paragraphs[0].text == "A"
    assert int(document.paragraphs[0].paragraph_format.left_indent) == int(docx._length(Dimen(18)))


def test_docx_paginated_hanging_label_collapses_leading_glue(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    para = pg.Paragraph(parser, indent=False)
    label = _FakeHBox(
        [
            nd.Glue(Glue(Dimen(-18)), None),
            *[nd.CharNode(ch, font) for ch in "[1]"],
            nd.Glue(Glue(Dimen(6)), None),
        ],
        width=0,
        rightmost_value=0,
    )
    line = _FakeHBox(
        [
            nd.Glue(Glue(Dimen(18)), "\\parindent"),
            label,
            nd.CharNode("A", font),
        ],
        para,
    )

    backend.shipout(_page_box([line]))

    document = _word_document(parser, backend)
    assert document.paragraphs[0].text == "[1] A"
    assert document.paragraphs[0].paragraph_format.left_indent is None


def test_docx_emits_destination_bookmark_for_pdf_dest_special(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    para = pg.Paragraph(parser, indent=False)
    line = _FakeHBox(
        [
            nd.Special("pdf: dest (target.1)[@thispage/XYZ @xpos @ypos null]"),
            nd.CharNode("A", font),
        ],
        para,
    )

    backend.shipout(_page_box([line]))

    root = _document_root(_docx_bytes(parser, backend))
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    bookmark = root.find(".//w:bookmarkStart", ns)
    assert bookmark is not None
    assert bookmark.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name") == "target.1"
    assert root.find(".//w:bookmarkEnd", ns) is not None


def test_docx_wraps_internal_goto_annotation_as_hyperlink(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    para = pg.Paragraph(parser, indent=False)
    line = _FakeHBox(
        [
            nd.Special("pdf: beginann <</Type/Annot/Subtype/Link/A<</S/GoTo/D(target.1)>>>>"),
            nd.CharNode("A", font),
            nd.Special("pdf: endann"),
            nd.Glue(Glue(Dimen(4)), None),
            nd.CharNode("B", font),
        ],
        para,
    )

    backend.shipout(_page_box([line]))

    root = _document_root(_docx_bytes(parser, backend))
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    hyperlink = root.find(".//w:hyperlink", ns)
    assert hyperlink is not None
    assert hyperlink.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}anchor") == "target.1"
    assert "".join(t.text or "" for t in hyperlink.findall(".//w:t", ns)) == "A"
    assert "".join(t.text or "" for t in root.findall(".//w:t", ns)) == "A B"


def test_docx_wraps_uri_annotation_as_external_hyperlink(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    para = pg.Paragraph(parser, indent=False)
    line = _FakeHBox(
        [
            nd.Special("pdf: beginann <</Type/Annot/Subtype/Link/A<</S/URI/URI(https://example.test/path)>>>>"),
            nd.CharNode("A", font),
            nd.Special("pdf: endann"),
        ],
        para,
    )

    backend.shipout(_page_box([line]))

    data = _docx_bytes(parser, backend)
    root = _document_root(data)
    rels = _document_relationships_root(data)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    hyperlink = root.find(".//w:hyperlink", ns)
    rid = hyperlink.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
    rel = next(
        node for node in rels
        if node.get("Id") == rid
    )
    assert rel.get("Target") == "https://example.test/path"
    assert rel.get("TargetMode") == "External"


def test_docx_reopens_line_spanning_annotation_per_line(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    para = pg.Paragraph(parser, indent=False)
    line1 = _FakeHBox(
        [
            nd.Special("pdf: beginann <</Type/Annot/Subtype/Link/A<</S/GoTo/D(target.1)>>>>"),
            nd.CharNode("A", font),
        ],
        para,
    )
    line2 = _FakeHBox(
        [
            nd.CharNode("B", font),
            nd.Special("pdf: endann"),
            nd.CharNode("C", font),
        ],
        para,
    )

    backend.shipout(_page_box([line1, line2]))

    root = _document_root(_docx_bytes(parser, backend))
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    hyperlinks = root.findall(".//w:hyperlink", ns)
    assert [
        link.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}anchor")
        for link in hyperlinks
    ] == ["target.1", "target.1"]
    assert ["".join(t.text or "" for t in link.findall(".//w:t", ns)) for link in hyperlinks] == ["A", "B"]
    assert "".join(t.text or "" for t in root.findall(".//w:t", ns)) == "ABC"


def test_docx_inline_vbox_emits_word_textbox_story(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    backend_baseline = Dimen(1)
    font = _install_font(parser, _FakeWordBaselineFont(backend_baseline))
    outer_para = pg.Paragraph(parser, indent=False)
    inner_para = pg.Paragraph(parser, indent=False)
    inner_line = _line_box("X", inner_para, font, width=16)
    vbox = _FakeVBox([inner_line], width=18, height=9, depth=3)
    line = _FakeHBox(
        [nd.CharNode("A", font), vbox, nd.CharNode("B", font)],
        outer_para,
    )

    backend.shipout(_page_box([line]))

    xml = _document_xml(_docx_bytes(parser, backend))
    assert "<w:drawing" in xml
    assert "<wp:inline" in xml
    assert "<w:txbxContent>" in xml
    assert "<w:t>A</w:t>" in xml
    assert "<w:t>X</w:t>" in xml
    assert "<w:t>B</w:t>" in xml
    wp_extent = re.search(r'<wp:extent\b[^>]*\bcx="([^"]+)"[^>]*\bcy="([^"]+)"', xml)
    effect = re.search(r'<wp:effectExtent\b[^>]*\bb="([^"]+)"', xml)
    shape_extent = re.search(r'<a:xfrm><a:off x="0" y="0"/><a:ext cx="([^"]+)" cy="([^"]+)"/>', xml)
    assert wp_extent.groups() == (
        str(docx._emu(vbox.width)),
        str(docx._emu(vbox.height + vbox.depth)),
    )
    assert effect.group(1) == "0"
    assert shape_extent.groups() == (
        str(docx._emu(vbox.width)),
        str(docx._emu(vbox.height + vbox.depth)),
    )
    expected_position = _box_position_to_tex_baseline(line.depth, vbox.depth, backend_baseline)
    assert f'<w:position w:val="{expected_position}"/>' in xml


def test_docx_epdf_graphic_special_converts_to_svg_picture(parser, monkeypatch):
    class FakeConverter:
        def convert(self, request):
            return graphics.GraphicAsset(
                format="svg",
                data="<svg xmlns='http://www.w3.org/2000/svg'></svg>",
                width=request.width,
                height=request.height,
                depth=request.depth,
            )

    monkeypatch.setitem(graphics._CONVERTERS, ("pdf", "svg"), FakeConverter())
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    _install_font(parser)
    source = pg.Paragraph(parser, indent=False)
    special = nd.Special("pdf: epdf bbox 0 0 200 100 width 72pt (fig.pdf)")

    backend.shipout(_page_box([_FakeHBox([special], source=source, width=72, height=36, depth=0)]))
    data = _docx_bytes(parser, backend)

    xml = _document_xml(data)
    assert "<asvg:svgBlip" in xml
    assert "<a14:useLocalDpi" in xml
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        media = [name for name in zf.namelist() if name.startswith("word/media/") and name.endswith(".png")]
        svg_media = [name for name in zf.namelist() if name.startswith("word/media/") and name.endswith(".svg")]
        assert len(media) == 1
        assert len(svg_media) == 1
        assert zf.read(media[0]).startswith(b"\x89PNG")


def test_docx_standalone_graphic_line_keeps_normal_edge_spacing(parser, monkeypatch):
    class FakeConverter:
        def convert(self, request):
            return graphics.GraphicAsset(
                format="svg",
                data="<svg xmlns='http://www.w3.org/2000/svg'></svg>",
                width=request.width,
                height=request.height,
                depth=request.depth,
            )

    monkeypatch.setitem(graphics._CONVERTERS, ("pdf", "svg"), FakeConverter())
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    _install_font(parser)
    source = pg.Paragraph(parser, indent=False)
    special = nd.Special("pdf: epdf bbox 0 0 200 100 width 72pt (fig.pdf)")
    graphic_box = _FakeHBox([special], width=72, height=36, depth=0)
    left = Dimen(12)
    right = Dimen(18)

    backend.shipout(
        _page_box(
            [
                _FakeHBox(
                    [nd.Glue(Glue(left), None), graphic_box, nd.Glue(Glue(right), None)],
                    source=source,
                    width=left + Dimen(72) + right,
                    height=36,
                    depth=0,
                )
            ]
        )
    )
    data = _docx_bytes(parser, backend)
    document = WordDocument(io.BytesIO(data))
    para = document.paragraphs[0]

    assert int(para.paragraph_format.left_indent) == int(docx._length(left))
    assert para.paragraph_format.right_indent is None
    assert para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    xml = _document_xml(data)
    assert "<w:drawing" in xml
    first_spacing = re.search(r"<w:p><w:pPr><w:spacing([^>]*)/>", xml)
    assert first_spacing is not None
    assert 'w:after="0"' in first_spacing.group(1)
    assert 'w:lineRule="exact"' in first_spacing.group(1)
    assert f'w:line="{docx._twips(Dimen(36))}"' in first_spacing.group(1)
    wp_extent = re.search(r'<wp:extent\b[^>]*\bcy="([^"]+)"', xml)
    assert wp_extent is not None
    assert int(wp_extent.group(1)) == docx._twip_emu(Dimen(36))


def test_docx_standalone_graphic_line_keeps_tex_line_height(parser, monkeypatch):
    class FakeConverter:
        def convert(self, request):
            return graphics.GraphicAsset(
                format="svg",
                data="<svg xmlns='http://www.w3.org/2000/svg'></svg>",
                width=request.width,
                height=request.height,
                depth=request.depth,
            )

    monkeypatch.setitem(graphics._CONVERTERS, ("pdf", "svg"), FakeConverter())
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    _install_font(parser)
    source = pg.Paragraph(parser, indent=False)
    special = nd.Special("pdf: epdf bbox 0 0 200 100 width 72pt (fig.pdf)")
    graphic_box = _FakeHBox([special], width=72, height=36, depth=0)
    line = _FakeHBox([graphic_box], source=source, width=72, height=40, depth=5)

    backend.shipout(_page_box([line]))

    xml = _document_xml(_docx_bytes(parser, backend))
    first_spacing = re.search(r"<w:p><w:pPr><w:spacing([^>]*)/>", xml)
    assert first_spacing is not None
    assert f'w:line="{docx._twips(Dimen(45))}"' in first_spacing.group(1)
    assert f'w:line="{docx._twips(Dimen(36))}"' not in first_spacing.group(1)


def test_docx_graphic_transform_advance_cancels_tex_kern(parser, monkeypatch):
    class FakeConverter:
        def convert(self, request):
            return graphics.GraphicAsset(
                format="svg",
                data="<svg xmlns='http://www.w3.org/2000/svg'></svg>",
                width=request.width,
                height=request.height,
                depth=request.depth,
            )

    monkeypatch.setitem(graphics._CONVERTERS, ("pdf", "svg"), FakeConverter())
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    _install_font(parser)
    source = pg.Paragraph(parser, indent=False)
    special = nd.Special("pdf: epdf bbox 0 0 200 100 width 72pt (fig.pdf)")
    graphic_box = _FakeHBox([special], width=72, height=36, depth=0)
    zero_width_graphic = _FakeHBox(
        [graphic_box, nd.Kern(Dimen(-72))],
        width=0,
        height=36,
        depth=0,
    )
    line = _FakeHBox(
        [
            nd.Special("pdf:btrans"),
            nd.Special("x:scale 0.5 0.5"),
            zero_width_graphic,
            nd.Special("pdf:etrans"),
            nd.Kern(Dimen(36)),
        ],
        source=source,
        width=36,
        height=18,
        depth=0,
    )

    backend.shipout(_page_box([line]))

    xml = _document_xml(_docx_bytes(parser, backend))
    drawing_runs = _drawing_runs(xml)
    assert len(drawing_runs) == 1
    assert re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml) == [""]
    assert '<w:spacing w:val=' not in xml


def test_docx_graphic_line_keeps_explicit_trailing_kern(parser, monkeypatch):
    class FakeConverter:
        def convert(self, request):
            return graphics.GraphicAsset(
                format="svg",
                data="<svg xmlns='http://www.w3.org/2000/svg'></svg>",
                width=request.width,
                height=request.height,
                depth=request.depth,
            )

    monkeypatch.setitem(graphics._CONVERTERS, ("pdf", "svg"), FakeConverter())
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    _install_font(parser)
    source = pg.Paragraph(parser, indent=False)
    special = nd.Special("pdf: epdf bbox 0 0 200 100 width 72pt (fig.pdf)")
    graphic_box = _FakeHBox([special], width=72, height=36, depth=0)
    line = _FakeHBox(
        [graphic_box, nd.Kern(Dimen(8))],
        source=source,
        width=80,
        height=36,
        depth=0,
    )

    backend.shipout(_page_box([line]))

    xml = _document_xml(_docx_bytes(parser, backend))
    assert len(_drawing_runs(xml)) == 1
    assert re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml) == ["", "\xa0"]
    assert '<w:spacing w:val=' in xml


def test_docx_graphic_followed_by_text_keeps_normal_space(parser, monkeypatch):
    class FakeConverter:
        def convert(self, request):
            return graphics.GraphicAsset(
                format="svg",
                data="<svg xmlns='http://www.w3.org/2000/svg'></svg>",
                width=request.width,
                height=request.height,
                depth=request.depth,
            )

    monkeypatch.setitem(graphics._CONVERTERS, ("pdf", "svg"), FakeConverter())
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    source = pg.Paragraph(parser, indent=False)
    special = nd.Special("pdf: epdf bbox 0 0 200 100 width 72pt (fig.pdf)")
    graphic_box = _FakeHBox([special], width=72, height=36, depth=0)

    backend.shipout(
        _page_box(
            [
                _FakeHBox(
                    [graphic_box, nd.Glue(Glue(Dimen(5)), None), nd.CharNode("A", font)],
                    source=source,
                    width=90,
                    height=36,
                    depth=0,
                )
            ]
        )
    )
    document = _word_document(parser, backend)

    assert document.paragraphs[0].text == " A"
    assert document.paragraphs[0].paragraph_format.right_indent is None


def test_docx_trailing_negative_spacing_keeps_tex_line_height(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    source = pg.Paragraph(parser, indent=False)
    line = _FakeHBox([nd.CharNode("A", font)], source=source, width=10, height=8, depth=2)

    backend.shipout(_page_box([line, nd.Kern(Dimen(-4))]))

    xml = _document_xml(_docx_bytes(parser, backend))
    first_spacing = re.search(r"<w:p><w:pPr><w:spacing([^>]*)/>", xml)
    assert first_spacing is not None
    assert f'w:line="{docx._twips(Dimen(10))}"' in first_spacing.group(1)
    assert f'w:line="{docx._twips(Dimen(6))}"' not in first_spacing.group(1)
    root = _document_root(_docx_bytes(parser, backend))
    pg_mar = root.find(f".//{{{docx._W_NS}}}pgMar")
    assert pg_mar is not None
    assert int(pg_mar.get(f"{{{docx._W_NS}}}bottom")) == (
        docx._twips(Dimen(72.27))
        - docx._twips(Dimen(4))
    )


def test_docx_graphic_inside_shifted_hbox_uses_parent_baseline(parser, monkeypatch):
    class FakeConverter:
        def convert(self, request):
            return graphics.GraphicAsset(
                format="svg",
                data="<svg xmlns='http://www.w3.org/2000/svg'></svg>",
                width=request.width,
                height=request.height,
                depth=request.depth,
            )

    monkeypatch.setitem(graphics._CONVERTERS, ("pdf", "svg"), FakeConverter())
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    backend_baseline = Dimen(2)
    _install_font(parser, _FakeWordBaselineFont(backend_baseline))
    source = pg.Paragraph(parser, indent=False)
    special = nd.Special("pdf: epdf bbox 0 0 200 100 width 72pt (fig.pdf)")
    graphic_box = _FakeHBox([special], width=72, height=36, depth=0)
    graphic_box.shifted = Dimen(4)
    line = _FakeHBox([graphic_box], source=source, width=72, height=36, depth=7)

    backend.shipout(_page_box([line]))

    xml = _document_xml(_docx_bytes(parser, backend))
    expected_position = _box_position_to_tex_baseline(
        line.depth,
        graphic_box.shifted,
        backend_baseline,
    )
    drawing_runs = _drawing_runs(xml)
    assert len(drawing_runs) == 1
    assert f'<w:position w:val="{expected_position}"/>' in drawing_runs[0]


def test_docx_graphic_line_moves_from_backend_to_tex_baseline(parser, monkeypatch):
    class FakeConverter:
        def convert(self, request):
            return graphics.GraphicAsset(
                format="svg",
                data="<svg xmlns='http://www.w3.org/2000/svg'></svg>",
                width=request.width,
                height=request.height,
                depth=request.depth,
            )

    monkeypatch.setitem(graphics._CONVERTERS, ("pdf", "svg"), FakeConverter())
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    backend_baseline = Dimen(7)
    font = _install_font(parser, _FakeWordBaselineFont(backend_baseline))
    source = pg.Paragraph(parser, indent=False)
    special = nd.Special("pdf: epdf bbox 0 0 200 100 width 72pt (fig.pdf)")
    graphic_box = _FakeHBox([special], width=72, height=36, depth=0)
    line = _FakeHBox([graphic_box], source=source, width=72, height=36, depth=0)

    backend.shipout(_page_box([line]))

    xml = _document_xml(_docx_bytes(parser, backend))
    expected_position = _box_position_to_tex_baseline(line.depth, Dimen(), backend_baseline)
    drawing_runs = _drawing_runs(xml)
    assert len(drawing_runs) == 1
    assert f'<w:position w:val="{expected_position}"/>' in drawing_runs[0]


def test_docx_inline_vbox_preserves_local_center_alignment(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    outer_para = pg.Paragraph(parser, indent=False)
    inner_para = pg.Paragraph(parser, indent=False)
    fill = Stretchness(Dimen(1), 1)
    inner_line = _FakeHBox(
        [
            nd.Glue(Glue(Dimen(), fill), None),
            nd.CharNode("A", font),
            nd.Glue(Glue(Dimen(), fill), None),
        ],
        inner_para,
        natural=Glue(Dimen(5), Stretchness(Dimen(2), 1)),
        glue_ratio=(1, int(Dimen(20)), int(Dimen(2))),
    )
    vbox = _FakeVBox([inner_line], width=80, height=9, depth=3)
    line = _FakeHBox([vbox], outer_para)

    backend.shipout(_page_box([line]))

    data = _docx_bytes(parser, backend)
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    word_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    textbox_paragraphs = root.findall(".//w:txbxContent//w:p", ns)
    jc_values = [
        jc.get(f"{word_ns}val")
        for p in textbox_paragraphs
        for jc in [p.find("w:pPr/w:jc", ns)]
        if jc is not None
    ]
    assert "center" in jc_values


def test_docx_inline_vbox_table_uses_exact_tex_widths(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    outer_para = pg.Paragraph(parser, indent=False)
    owner = align.HAlignment()
    owner.tabskips = [
        Glue(Dimen(2), Stretchness(Dimen(1), 0)),
        Glue(Dimen(3), Stretchness(Dimen(3), 0)),
        Glue(Dimen(5)),
    ]
    row = align.Row()
    row.cells = [_alignment_cell("a", font, width=10), _alignment_cell("b", font, width=20)]
    owner.rows = [row]
    natural = Glue(Dimen(40), Stretchness(Dimen(4), 0))
    row_box = _FakeHBox(
        [],
        source=owner,
        width=48,
        height=8,
        depth=2,
        natural=natural,
        glue_ratio=(1, int(Dimen(8)), int(Dimen(4))),
    )
    vbox = _FakeVBox(
        [row_box],
        width=48,
        height=10,
        depth=0,
    )
    line = _FakeHBox(
        [nd.CharNode("A", font), vbox, nd.CharNode("B", font)],
        outer_para,
    )

    backend.shipout(_page_box([line]))

    xml = _document_xml(_docx_bytes(parser, backend))
    table_width = re.search(r'<w:tblW\b[^>]*\bw:type="([^"]+)"[^>]*\bw:w="([^"]+)"', xml)
    assert table_width.groups() == ("dxa", docx.twips(vbox.width))
    grid_widths = [
        int(value)
        for value in re.findall(r'<w:gridCol\b[^>]*\bw:w="(\d+)"', xml)
    ]
    assert grid_widths == [
        int(docx.twips(Dimen(4))),
        int(docx.twips(Dimen(10))),
        int(docx.twips(Dimen(9))),
        int(docx.twips(Dimen(20))),
        int(docx.twips(Dimen(5))),
    ]
    textbox_xml = re.search(r"<w:txbxContent>(.*)</w:txbxContent>", xml, re.DOTALL).group(1)
    line_heights = re.findall(r'<w:spacing\b[^>]*\bw:line="(\d+)"', textbox_xml)
    assert line_heights == [docx.twips(row_box.height + row_box.depth)] * 2
    extent_width = int(re.search(r"<wp:extent[^>]* cx=\"(\d+)\"", xml).group(1))
    assert extent_width == docx._emu(vbox.width)
    assert 1440 not in grid_widths


def test_docx_explicit_glue_emits_preserved_space_with_spacing_hint(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    backend_baseline = Dimen(1)
    font = _install_font(parser, _FakeWordBaselineFont(backend_baseline))
    para = pg.Paragraph(parser, indent=False)
    line = _FakeHBox(
        [
            nd.CharNode("A", font),
            nd.Glue(Glue(Dimen(9)), None),
            nd.CharNode("B", font),
        ],
        para,
    )

    backend.shipout(_page_box([line]))

    xml = _document_xml(_docx_bytes(parser, backend))
    assert 'xml:space="preserve"' in xml
    assert f'w:val="{docx.twips(Dimen(4))}"' in xml
    expected_position = _text_position_to_tex_baseline(line.depth, backend_baseline)
    assert re.findall(r'<w:position w:val="([^"]+)"/>', xml) == [expected_position] * 3


def test_docx_spacing_uses_scaled_font_space_width(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser, _FakeFont(size=20))
    para = pg.Paragraph(parser, indent=False)
    line = _FakeHBox(
        [
            nd.CharNode("A", font),
            nd.Glue(Glue(Dimen(12)), None),
            nd.CharNode("B", font),
        ],
        para,
    )

    backend.shipout(_page_box([line]))

    xml = _document_xml(_docx_bytes(parser, backend))
    assert f'w:val="{docx.twips(Dimen(2))}"' in xml
    assert f'w:val="{docx.twips(Dimen(7))}"' not in xml


def test_docx_negative_space_spacing_rounds_down(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    para = pg.Paragraph(parser, indent=False)
    width = Dimen(5) + docx._tex_points(-0.099)
    line = _FakeHBox(
        [
            nd.CharNode("A", font),
            nd.Glue(Glue(width), None),
            nd.CharNode("B", font),
        ],
        para,
    )

    backend.shipout(_page_box([line]))

    xml = _document_xml(_docx_bytes(parser, backend))
    assert 'w:val="-2"' in xml


def test_docx_centered_hbox_uses_tex_glue_not_word_alignment(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    para = pg.Paragraph(parser, indent=False)
    fill = Stretchness(Dimen(1), 1)
    left = nd.Glue(Glue(Dimen(), fill), None)
    right = nd.Glue(Glue(Dimen(), fill), None)
    natural = Glue(Dimen(5), Stretchness(Dimen(2), 1))
    line = _FakeHBox(
        [left, nd.CharNode("A", font), right],
        para,
        natural=natural,
        glue_ratio=(1, int(Dimen(20)), int(Dimen(2))),
    )

    backend.shipout(_page_box([line]))

    document = _word_document(parser, backend)
    assert document.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert int(document.paragraphs[0].paragraph_format.left_indent) == int(docx._length(Dimen(10)))


def test_docx_inline_math_embeds_svg_picture(parser, monkeypatch):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    para = pg.Paragraph(parser, indent=False)
    owner = mmode.InlineMathNode(nodes=[_math_atom("x")])
    on = nd.MathShift(True)
    on.source = owner
    on.kern = Dimen()
    off = nd.MathShift(False)
    off.source = owner
    off.kern = Dimen()
    line = _FakeHBox(
        [nd.CharNode("A", font), on, nd.CharNode("x", font), off, nd.CharNode("B", font)],
        para,
        height=12,
        depth=4,
    )
    captured = {}

    def fake_svg(box):
        captured["box"] = box
        return b'<svg xmlns="http://www.w3.org/2000/svg" width="1" height="1"/>'

    monkeypatch.setattr(backend, "inlineMathSvg", fake_svg)

    backend.shipout(_page_box([line]))

    document = _word_document(parser, backend)
    assert document.paragraphs[0].text == "AB"
    data = _docx_bytes(parser, backend)
    xml = _document_xml(data)
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        names = set(zf.namelist())
        rels_xml = zf.read("word/_rels/document.xml.rels").decode("utf-8")
        content_types = zf.read("[Content_Types].xml").decode("utf-8")
        svg_payload = zf.read("word/media/pytex-inline-math-1.svg")
    assert captured["box"].list[0].char == "x"
    assert "word/media/pytex-inline-math-1.svg" in names
    assert svg_payload == b'<svg xmlns="http://www.w3.org/2000/svg" width="1" height="1"/>'
    assert "<w:drawing" in xml
    assert "<pic:pic>" in xml
    assert "<asvg:svgBlip" in xml
    assert "<a:noFill/>" in xml
    assert "pytexInlineSvg" not in xml
    wp_extent = re.search(r'<wp:extent\b[^>]*\bcx="([^"]+)"[^>]*\bcy="([^"]+)"', xml)
    effect = re.search(r'<wp:effectExtent\b[^>]*\bt="([^"]+)"[^>]*\bb="([^"]+)"', xml)
    transform = re.search(
        r'<a:xfrm><a:off x="0" y="([^"]+)"/><a:ext cx="([^"]+)" cy="([^"]+)"/>',
        xml,
    )
    assert wp_extent.groups() == (
        str(docx._emu(captured["box"].width)),
        str(docx._emu(captured["box"].height + captured["box"].depth)),
    )
    assert effect.groups() == ("0", "0")
    assert transform.groups() == (
        "0",
        str(docx._emu(captured["box"].width)),
        str(docx._emu(captured["box"].height + captured["box"].depth)),
    )
    drawing_runs = _drawing_runs(xml)
    assert len(drawing_runs) == 1
    assert "<w:position" not in drawing_runs[0]
    assert 'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"' in rels_xml
    assert 'Target="media/pytex-inline-math-1.svg"' in rels_xml
    assert 'Extension="svg"' in content_types
    assert "image/svg+xml" in content_types


def test_docx_inline_math_alignment_vbox_uses_current_text_run(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    para = pg.Paragraph(parser, indent=False)
    alignment = align.HAlignment()
    alignment.tabskips = [Glue(Dimen(2)), Glue(Dimen(3))]
    row = align.Row()
    row.cells = [_alignment_cell("x", font, width=10)]
    alignment.rows = [row]
    row_box = _alignment_row_box(alignment, row, height=8, depth=2)
    vbox = bx.VBox(parser, None, 0)
    vbox.list = [row_box]
    vbox.width = Dimen(15)
    vbox.height = Dimen(10)
    vbox.depth = Dimen()
    vbox.to = vbox.height
    vbox.spread = Dimen()
    vbox.natural = Glue()
    vbox.glue_ratio = bx.GlueRatio(0, 0, 1)
    vbox.shifted = Dimen()
    owner = mmode.InlineMathNode(nodes=[mmode.Box(vbox)])
    on = nd.MathShift(True)
    on.source = owner
    on.kern = Dimen()
    off = nd.MathShift(False)
    off.source = owner
    off.kern = Dimen()
    line = _FakeHBox(
        [nd.CharNode("A", font), on, off, nd.CharNode("B", font)],
        para,
        height=10,
        depth=2,
    )

    backend.shipout(_page_box([line]))

    xml = _document_xml(_docx_bytes(parser, backend))
    assert "<w:txbxContent>" in xml
    assert "<w:tbl>" in xml
    assert "<w:t>A</w:t>" in xml
    assert "<w:t>x</w:t>" in xml
    assert "<w:t>B</w:t>" in xml


def test_docx_inline_svg_placeholders_do_not_collide_after_ninth_picture(parser, monkeypatch):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    para = pg.Paragraph(parser, indent=False)
    nodes = []
    for _ in range(12):
        owner = mmode.InlineMathNode(nodes=[_math_atom("x")])
        on = nd.MathShift(True)
        on.source = owner
        on.kern = Dimen()
        off = nd.MathShift(False)
        off.source = owner
        off.kern = Dimen()
        nodes.extend([on, nd.CharNode("x", font), off, nd.Glue(Glue(Dimen(2)), None)])
    line = _FakeHBox(nodes, para, width=120, height=12, depth=4)

    def fake_svg(box):
        return b'<svg xmlns="http://www.w3.org/2000/svg" width="1" height="1"/>'

    monkeypatch.setattr(backend, "inlineMathSvg", fake_svg)

    backend.shipout(_page_box([line]))

    data = _docx_bytes(parser, backend)
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        names = set(zf.namelist())
        document_xml = zf.read("word/document.xml").decode("utf-8")
        rels_xml = zf.read("word/_rels/document.xml.rels").decode("utf-8")
    embed_ids = set(re.findall(r'r:embed="([^"]+)"', document_xml))
    image_rel_ids = set(
        re.findall(
            r'<Relationship Id="([^"]+)" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"',
            rels_xml,
        )
    )
    assert "pytexInlineSvg" not in document_xml
    assert len(embed_ids) == 24
    assert embed_ids <= image_rel_ids
    assert len([name for name in names if name.startswith("word/media/") and name.endswith(".svg")]) == 12
    assert len([name for name in names if name.startswith("word/media/") and name.endswith(".png")]) == 12


def test_docx_inline_math_inside_shifted_hbox_uses_parent_baseline(parser, monkeypatch):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    backend_baseline = Dimen(2)
    font = _install_font(parser, _FakeWordBaselineFont(backend_baseline))
    para = pg.Paragraph(parser, indent=False)
    owner = mmode.InlineMathNode(nodes=[_math_atom("x")])
    on = nd.MathShift(True)
    on.source = owner
    on.kern = Dimen()
    off = nd.MathShift(False)
    off.source = owner
    off.kern = Dimen()
    inner = _FakeHBox([on, nd.CharNode("x", font), off], width=10, height=8, depth=3)
    inner.shifted = Dimen(2)
    line = _FakeHBox([inner], para, width=10, height=6, depth=5)

    captured = {}

    def fake_svg(box):
        captured["box"] = box
        return b'<svg xmlns="http://www.w3.org/2000/svg" width="1" height="1"/>'

    monkeypatch.setattr(backend, "inlineMathSvg", fake_svg)

    backend.shipout(_page_box([line]))

    xml = _document_xml(_docx_bytes(parser, backend))
    expected_position = _box_position_to_tex_baseline(
        line.depth,
        captured["box"].depth + inner.shifted,
        backend_baseline,
    )
    drawing_runs = _drawing_runs(xml)
    assert len(drawing_runs) == 1
    assert f'<w:position w:val="{expected_position}"/>' in drawing_runs[0]


def test_docx_display_math_embeds_shifted_svg_picture(parser, monkeypatch):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    backend_baseline = Dimen(2)
    font = _install_font(parser, _FakeWordBaselineFont(backend_baseline))
    owner = mmode.DisplayMathNode()
    owner.list.append(_math_atom("x"))
    before = nd.Glue(Glue(Dimen(6)), "\\abovedisplayskip")
    before.source = owner
    after = nd.Glue(Glue(Dimen(4)), "\\belowdisplayskip")
    after.source = owner
    display_box = _FakeHBox([nd.CharNode("x", font)], source=owner, width=50, height=12, depth=3)
    display_box.shifted = Dimen(20)
    para = pg.Paragraph(parser, indent=False)
    captured = []

    def fake_svg(box):
        captured.append(box)
        return b'<svg xmlns="http://www.w3.org/2000/svg" width="1" height="1"/>'

    monkeypatch.setattr(backend, "inlineMathSvg", fake_svg)
    monkeypatch.setattr(docx, "_svg_png_fallback", lambda payload: b"\x89PNG\r\n\x1a\nfallback")

    backend.shipout(_page_box([before, display_box, after, _line_box("After", para, font)]))

    document = _word_document(parser, backend)
    assert document.paragraphs[0].text == ""
    assert document.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert document.paragraphs[1].text == "After"
    assert int(document.paragraphs[1].paragraph_format.space_before) == int(docx._length(Dimen(4)))
    assert len(captured) == 1
    assert captured[0].node_type == nd.NODE_TYPE.VLIST
    assert captured[0].list == [display_box]
    assert captured[0].width == display_box.shifted + display_box.width
    data = _docx_bytes(parser, backend)
    xml = _document_xml(data)
    assert "<w:drawing" in xml
    assert "<asvg:svgBlip" not in xml
    assert f'cx="{docx._emu(display_box.shifted + display_box.width)}"' in xml
    drawing_runs = _drawing_runs(xml)
    assert len(drawing_runs) == 1
    expected_position = _box_position_to_tex_baseline(
        display_box.depth,
        display_box.depth,
        backend_baseline,
    )
    assert f'<w:position w:val="{expected_position}"/>' in drawing_runs[0]
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        names = set(zf.namelist())
        rels_xml = zf.read("word/_rels/document.xml.rels").decode("utf-8")
        assert "word/media/pytex-inline-math-1.png" in names
        assert "word/media/pytex-inline-math-1.svg" not in names
        assert zf.read("word/media/pytex-inline-math-1.png") == b"\x89PNG\r\n\x1a\nfallback"
        assert 'Target="media/pytex-inline-math-1.png"' in rels_xml
        assert 'Target="media/pytex-inline-math-1.svg"' not in rels_xml


def test_docx_display_math_uses_page_glue_state_for_display_skip(parser, monkeypatch):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    owner = mmode.DisplayMathNode()
    before = nd.Glue(
        Glue(Dimen(12), shrink=Stretchness(Dimen(12), 0)),
        "\\abovedisplayskip",
    )
    before.source = owner
    display_box = _FakeHBox([nd.CharNode("x", font)], source=owner, width=50, height=12, depth=3)
    page = _FakeVBox([before, display_box], width=200, height=60)
    page.natural = Glue(Dimen(), shrink=Stretchness(Dimen(12), 0))
    page.glue_ratio = bx.GlueRatio(-1, int(Dimen(6)), int(Dimen(12)))

    def fake_svg(box):
        return b'<svg xmlns="http://www.w3.org/2000/svg" width="1" height="1"/>'

    monkeypatch.setattr(backend, "inlineMathSvg", fake_svg)

    backend.shipout(page)

    document = _word_document(parser, backend)
    assert len(document.paragraphs) == 1
    assert int(document.paragraphs[0].paragraph_format.space_before) == int(docx._length(Dimen(6)))


def test_docx_vlist_tail_negative_glue_keeps_last_line_layout(parser, monkeypatch):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    owner = mmode.DisplayMathNode()
    before = nd.Glue(Glue(Dimen(6)), "\\abovedisplayskip")
    before.source = owner
    after = nd.Glue(Glue(Dimen(-4)), None)
    display_box = _FakeHBox([nd.CharNode("x", font)], source=owner, width=50, height=12, depth=6)

    def fake_svg(box):
        return b'<svg xmlns="http://www.w3.org/2000/svg" width="1" height="1"/>'

    monkeypatch.setattr(backend, "inlineMathSvg", fake_svg)

    backend.shipout(_page_box([before, display_box, after]))

    document = _word_document(parser, backend)
    assert len(document.paragraphs) == 1
    assert int(document.paragraphs[0].paragraph_format.line_spacing) == int(docx._length(Dimen(18)))
    xml = _document_xml(_docx_bytes(parser, backend))
    wp_extent = re.search(r'<wp:extent\b[^>]*\bcx="([^"]+)"[^>]*\bcy="([^"]+)"', xml)
    effect = re.search(r'<wp:effectExtent\b[^>]*\bt="([^"]+)"[^>]*\bb="([^"]+)"', xml)
    transform = re.search(
        r'<a:xfrm><a:off x="0" y="([^"]+)"/><a:ext cx="([^"]+)" cy="([^"]+)"/>',
        xml,
    )
    assert wp_extent.groups() == (
        str(docx._emu(display_box.width)),
        str(docx._emu(Dimen(18))),
    )
    assert effect.groups() == ("0", "0")
    assert transform.groups() == (
        "0",
        str(docx._emu(display_box.width)),
        str(docx._emu(Dimen(18))),
    )


def test_docx_vlist_tail_positive_glue_does_not_expand_last_line(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    para = pg.Paragraph(parser, indent=False)
    line = _FakeHBox(_char_nodes("Final", font), source=para, height=9, depth=3)

    backend.shipout(_page_box([line, nd.Glue(Glue(Dimen(200)), None)]))

    document = _word_document(parser, backend)
    assert len(document.paragraphs) == 1
    assert int(document.paragraphs[0].paragraph_format.line_spacing) == int(docx._length(Dimen(12)))


def test_docx_alignment_should_emit_word_table(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    owner = align.HAlignment()
    owner.tabskips = [Glue(Dimen(0)), Glue(Dimen(0)), Glue(Dimen(0))]
    row1 = align.Row()
    row1.cells = [_alignment_cell("a", font), _alignment_cell("b", font)]
    row2 = align.Row()
    row2.cells = [_alignment_cell("c", font), _alignment_cell("d", font)]
    owner.rows = [row1, row2]

    backend.shipout(_page_box([
        _alignment_row_box(owner, row1),
        _alignment_row_box(owner, row2),
    ]))

    data = _docx_bytes(parser, backend)
    document = WordDocument(io.BytesIO(data))
    assert len(document.tables) == 1
    assert document.tables[0].alignment == WD_TABLE_ALIGNMENT.LEFT
    assert document.tables[0].cell(0, 1).text == "a"
    assert document.tables[0].cell(0, 3).text == "b"
    assert document.tables[0].cell(1, 1).text == "c"
    assert document.tables[0].cell(1, 3).text == "d"
    xml = _document_xml(data)
    assert "<w:tblCellMar>" in xml
    assert xml.count('w:w="0" w:type="dxa"') >= 4
    assert xml.count("<w:noWrap/>") >= 4


def test_docx_alignment_zero_width_overhang_cell_gets_visible_width(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    owner = align.HAlignment()
    owner.tabskips = [Glue(Dimen(50)), Glue(Dimen(50))]
    visible = bx.HBox(parser, None, 0)
    visible.list = _char_nodes("1", font)
    visible = visible.typeset(parser)
    cell = bx.HBox(parser, Dimen(), None)
    cell.list = [nd.Kern(-visible.width), visible]
    cell = cell.typeset(parser)
    cell.span = 1
    row = align.Row()
    row.cells = [cell]
    owner.rows = [row]
    row_box = _alignment_row_box(owner, row, height=10, depth=2)

    backend.shipout(_page_box([row_box]))

    xml = _document_xml(_docx_bytes(parser, backend))
    cell_texts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', xml)
    grid_widths = [
        value
        for value in re.findall(r'<w:gridCol\b[^>]*\bw:w="(\d+)"', xml)
    ]
    assert grid_widths == [
        docx.twips(Dimen(50) - visible.width),
        docx.twips(visible.width),
        docx.twips(Dimen(50)),
    ]
    assert '<w:tblLayout w:type="fixed"/>' in xml
    assert ">1<" in xml
    assert "\xa0" not in "".join(cell_texts)


def test_docx_alignment_zero_width_right_protrusion_takes_following_width(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    owner = align.HAlignment()
    owner.tabskips = [Glue(Dimen(50)), Glue(Dimen(50))]
    visible = bx.HBox(parser, None, 0)
    visible.list = _char_nodes("1", font)
    visible = visible.typeset(parser)
    hss = Glue(Dimen(), Stretchness(Dimen(1), 1), Stretchness(Dimen(1), 1))
    cell = bx.HBox(parser, Dimen(), None)
    cell.list = [visible, nd.Glue(hss, None)]
    cell = cell.typeset(parser)
    cell.span = 1
    row = align.Row()
    row.cells = [cell]
    owner.rows = [row]

    backend.shipout(_page_box([_alignment_row_box(owner, row, height=10, depth=2)]))

    xml = _document_xml(_docx_bytes(parser, backend))
    grid_widths = [
        value
        for value in re.findall(r'<w:gridCol\b[^>]*\bw:w="(\d+)"', xml)
    ]
    assert grid_widths == [
        docx.twips(Dimen(50)),
        docx.twips(visible.width),
        docx.twips(Dimen(50) - visible.width),
    ]


def test_docx_alignment_cell_inline_math_uses_table_row_baseline(parser, monkeypatch):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    backend_baseline = Dimen(2)
    font = _install_font(parser, _FakeWordBaselineFont(backend_baseline))
    owner = align.HAlignment()
    math_owner = mmode.InlineMathNode(nodes=[_math_atom("x")])
    on = nd.MathShift(True)
    on.source = math_owner
    on.kern = Dimen()
    off = nd.MathShift(False)
    off.source = math_owner
    off.kern = Dimen()
    cell = _FakeHBox([on, nd.CharNode("x", font), off], width=10, height=8, depth=3)
    cell.span = 1
    row = align.Row()
    row.cells = [cell]
    owner.rows = [row]
    row_box = _alignment_row_box(owner, row, height=8, depth=3)

    captured = {}

    def fake_svg(box):
        captured["box"] = box
        return b'<svg xmlns="http://www.w3.org/2000/svg" width="1" height="1"/>'

    monkeypatch.setattr(backend, "inlineMathSvg", fake_svg)

    backend.shipout(_page_box([row_box]))

    xml = _document_xml(_docx_bytes(parser, backend))
    assert "<w:drawing" in xml
    expected_position = _box_position_to_tex_baseline(
        row_box.depth,
        captured["box"].depth,
        backend_baseline,
    )
    drawing_runs = _drawing_runs(xml)
    assert len(drawing_runs) == 1
    assert f'<w:position w:val="{expected_position}"/>' in drawing_runs[0]


def test_docx_zero_width_inline_vbox_is_not_emitted(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    para = pg.Paragraph(parser, indent=False)
    vbox = _FakeVBox([], width=0, height=8, depth=3)
    line = _FakeHBox([nd.CharNode("A", font), vbox, nd.CharNode("B", font)], para)

    backend.shipout(_page_box([line]))

    data = _docx_bytes(parser, backend)
    xml = _document_xml(data)
    document = WordDocument(io.BytesIO(data))
    assert "<w:drawing" not in xml
    assert document.paragraphs[0].text == "AB"


def test_docx_inline_vbox_moves_from_backend_to_tex_baseline(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    backend_baseline = Dimen(4)
    font = _install_font(parser, _FakeWordBaselineFont(backend_baseline))
    para = pg.Paragraph(parser, indent=False)
    char = nd.CharNode("A", font)
    vbox = _FakeVBox([], width=20, height=8, depth=3)
    line = _FakeHBox([char, vbox], para, height=8, depth=3)

    backend.shipout(_page_box([line]))

    xml = _document_xml(_docx_bytes(parser, backend))
    assert "<w:drawing" in xml
    expected_position = _box_position_to_tex_baseline(line.depth, vbox.depth, backend_baseline)
    drawing_runs = _drawing_runs(xml)
    assert len(drawing_runs) == 1
    assert f'<w:position w:val="{expected_position}"/>' in drawing_runs[0]


def test_docx_inline_vbox_inside_shifted_hbox_uses_parent_baseline(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    backend_baseline = Dimen(1)
    _install_font(parser, _FakeWordBaselineFont(backend_baseline))
    para = pg.Paragraph(parser, indent=False)
    vbox = _FakeVBox([], width=20, height=8, depth=3)
    shifted = _FakeHBox([vbox], width=20, height=8, depth=3)
    shifted.shifted = Dimen(2)
    line = _FakeHBox([shifted], para, width=20, height=6, depth=7)

    backend.shipout(_page_box([line]))

    xml = _document_xml(_docx_bytes(parser, backend))
    expected_position = _box_position_to_tex_baseline(
        line.depth,
        vbox.depth + shifted.shifted,
        backend_baseline,
    )
    drawing_runs = _drawing_runs(xml)
    assert len(drawing_runs) == 1
    assert f'<w:position w:val="{expected_position}"/>' in drawing_runs[0]


def test_docx_inline_vtop_uses_depth_baseline(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    backend_baseline = Dimen(2)
    _install_font(parser, _FakeWordBaselineFont(backend_baseline))
    para = pg.Paragraph(parser, indent=False)
    vtop = bx.VTop(parser, None, 0)
    vtop.list = []
    vtop.width = Dimen(20)
    vtop.height = Dimen(8)
    vtop.depth = Dimen(30)
    vtop.to = vtop.height
    vtop.spread = Dimen()
    vtop.natural = Glue()
    vtop.glue_ratio = bx.GlueRatio(0, 0, 1)
    vtop.shifted = Dimen()
    line = _FakeHBox([vtop], para, width=20, height=8, depth=30)

    backend.shipout(_page_box([line]))

    xml = _document_xml(_docx_bytes(parser, backend))
    assert "<w:drawing" in xml
    expected_position = _box_position_to_tex_baseline(line.depth, vtop.depth, backend_baseline)
    assert f'<w:position w:val="{expected_position}"/>' in xml


def test_docx_alignment_span_merges_word_cells(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    owner = align.HAlignment()
    owner.tabskips = []
    row = align.Row()
    cell = _alignment_cell("wide", font)
    cell.span = 2
    row.cells = [cell]
    owner.rows = [row]

    backend.shipout(_page_box([_alignment_row_box(owner, row)]))

    data = _docx_bytes(parser, backend)
    document = WordDocument(io.BytesIO(data))
    assert len(document.tables[0].columns) == 2
    assert document.tables[0].cell(0, 0).text == "wide"
    assert 'w:gridSpan w:val="2"' in _document_xml(data)


def test_docx_alignment_row_heights_include_tex_interline_spacing(parser):
    backend = docx.DocxBackend(parser)
    parser.shipout = backend
    font = _install_font(parser)
    owner = align.HAlignment()
    owner.tabskips = []
    row1 = align.Row()
    row1.cells = [_alignment_cell("a", font, width=15)]
    row2 = align.Row()
    row2.cells = [_alignment_cell("b", font, width=15)]
    owner.rows = [row1, row2]
    rowbox1 = _alignment_row_box(owner, row1, height=8, depth=2)
    rowbox2 = _alignment_row_box(owner, row2, height=7, depth=3)
    interline = nd.Glue(Glue(Dimen(5)), "\\baselineskip")
    interline.source = rowbox2

    backend.shipout(_FakeVBox([rowbox1, interline, rowbox2]))

    xml = _document_xml(_docx_bytes(parser, backend))
    heights = [
        int(value)
        for value in re.findall(r'<w:trHeight\b[^>]*\bw:val="(\d+)"', xml)
    ]
    assert heights == [int(docx.twips(Dimen(10))), int(docx.twips(Dimen(15)))]
    assert f'w:before="{docx.twips(Dimen(5))}"' in xml
